"""render, manifest, the ICD fetch, the identity stage, and the four gate
modules the registry names."""
from __future__ import annotations

import io
import json
import zipfile
from pathlib import Path

from medchron import config as config_mod, job as job_mod
from medchron.gates import claim_audit, cross_client, extractive, provenance
from medchron.stages import icd_fetch, identity, manifest, render
from medchron.stages.base import StageRun
from medchron_testkit import FakeSeat, job_yaml, make_pdf

DOC = """Alpha Example - Medical Chronology

## Treatment Timeline

| Medical Provider | Treatment Period | Visits | Reference |
| Example Clinic | 01/20/2026 - 03/30/2026 | 2 | Exhibit 1 - p. 2 |

## Diagnostic Highlights

| ICD Code | Description | First Diagnosed | Reference |
| S13.4XXA | Sprain of ligaments of cervical spine | 01/20/2026 | Exhibit 1 - p. 2 |

## Medical Chronology

Prior Medical History

**Neck:** chronic pain before the incident. (Exhibit 1 - p. 1)

[NTD: 1 additional pre-incident encounter was reviewed and is not itemized.]

01/20/2026
Example Clinic | Patient Complaints & Limitations

Neck pain after the collision. (Exhibit 1 - p. 2)

Medical Diagnoses

Cervical strain. (Exhibit 1 - p. 2)

03/30/2026
Example Clinic | Treatment Recommendations

Continue therapy. (Exhibit 1 - p. 3)

## Exhibit List

| Exhibit No. | Description |
| 1 | Example Clinic - 01-20-2026 - 03-30-2026 (Medical Records) |

## Records Reviewed and Limitations

This chronology was prepared from 2 documents in the matter file.

* "stray.pdf" (/MEDICAL) - could not be retrieved
"""


def _sr(job_dir: Path, firm: Path, data_root: Path, log: list[str] | None = None) -> StageRun:
    job = job_mod.load(job_dir)
    cfg = config_mod.load(str(firm))
    lines = log if log is not None else []
    return StageRun(job=job, cfg=cfg, unit=job.units[0], slug_dir=data_root / "example-matter", decided={},
                    log=lines.append, seat_factory=lambda: FakeSeat([], [], {}), date_stamp="08-29-26")


def test_render_writes_the_document_with_visuals_gap_bar_and_front_review_note(job_dir: Path, firm_config_path: Path, data_root: Path) -> None:
    from docx import Document

    log: list[str] = []
    sr = _sr(job_dir, firm_config_path, data_root, log)
    rd = sr.slug_dir / "runs" / "alpha"
    rd.mkdir(parents=True)
    (rd / "final-chronology.md").write_text(DOC)
    assert render.run(sr) == 0
    out = sr.slug_dir / "out" / "alpha"
    dst = out / "Alpha Example - Medical Chronology 08-29-26.docx"
    assert dst.is_file() and (out / "img" / "timeline.png").is_file() and (out / "img" / "cal-2026.png").is_file()
    d = Document(str(dst))
    paras = [p.text for p in d.paragraphs]
    assert paras[0] == "Alpha Example - Medical Chronology"
    assert paras[1] == "Records Reviewed and Limitations" and "stray.pdf" in " ".join(paras[:6])   # at the front
    assert any("Treatment Gap" in p and "69 days" in p for p in paras)                    # 01/20 -> 03/30 within the claim
    assert "Prior Medical History" in paras
    bold_runs = [r.text for p in d.paragraphs for r in p.runs if r.bold]
    assert "Neck:" in bold_runs and not any("**" in p for p in paras)
    cells = [c.text for t in d.tables for row in t.rows for c in row.cells]
    assert "Exhibit 1 - p. 2" in cells and "Cervical strain." in " ".join(cells)
    assert any("2 entries" in line and "1 gap bar" in line for line in log)


def test_manifest_names_every_deliverable_and_refuses_a_missing_one(job_dir: Path, firm_config_path: Path, data_root: Path) -> None:
    sr = _sr(job_dir, firm_config_path, data_root)
    out = sr.slug_dir / "out" / "alpha"
    out.mkdir(parents=True)
    (out / "Exhibit 2 - Example Imaging - 02-02-2026 (Medical Records).pdf").write_bytes(make_pdf([""]))
    (out / "Exhibit 1 - Example Clinic - 01-20-2026 (Medical Records).pdf").write_bytes(make_pdf([""]))
    assert manifest.run(sr) == 1                                             # no chronology yet
    (out / "Alpha Example - Medical Chronology 08-29-26.docx").write_bytes(b"docx")
    (out / "Alpha Example - Medical Billing Worksheet 08-29-26.docx").write_bytes(b"docx")
    assert manifest.run(sr) == 0
    m = json.loads((out / "upload_manifest.json").read_text())
    assert [x["name"].split(" - ")[0] for x in m] == ["Alpha Example", "Alpha Example", "Exhibit 1", "Exhibit 2"]
    assert m[0]["folder"] == "MEDICAL CHRONOLOGY - Alpha Example 08-29-26 by Example Operator" and m[0]["bytes"] == 4


def test_icd_fetch_vendors_the_tables_with_a_version_record(job_dir: Path, firm_config_path: Path, data_root: Path, tmp_path: Path) -> None:
    def zipped(member: str, body: bytes) -> bytes:
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr(f"folder/{member}", body)
        return buf.getvalue()

    blobs = {icd_fetch.ICD10_URL: zipped(icd_fetch.ICD10_MEMBER, b"00001 S134XXA 1 x y\n"),
             icd_fetch.ICD9_URL: zipped(icd_fetch.ICD9_MEMBER, b"7242 Lumbago\n")}
    v = icd_fetch.vendor(tmp_path / "icd", fetch=lambda url: blobs[url])
    assert (tmp_path / "icd" / "icd10cm_order.txt").read_bytes() == b"00001 S134XXA 1 x y\n"
    assert v["icd9cm"]["sha256"] and json.loads((tmp_path / "icd" / "VERSION.json").read_text())["icd10cm"]["url"] == icd_fetch.ICD10_URL
    sr = _sr(job_dir, firm_config_path, data_root)
    (sr.job.data_root / "controls" / "icd").mkdir(parents=True)
    (sr.job.data_root / "controls" / "icd" / "VERSION.json").write_text("{}")
    assert icd_fetch.run(sr) == 0                                             # present: no fetch


def test_cross_client_gate_flags_a_file_naming_the_other_client(tmp_path: Path, firm_config_path: Path, data_root: Path) -> None:
    jd = tmp_path / "job"
    jd.mkdir()
    (jd / "job.yaml").write_text(job_yaml(data_root, joint=True))
    log: list[str] = []
    sr = _sr(jd, firm_config_path, data_root, log)
    d = sr.slug_dir
    (d / "units").mkdir(parents=True)
    (d / "text").mkdir()
    (d / "runs" / "alpha").mkdir(parents=True)
    (d / "units.json").write_text(json.dumps({"alpha": {"surname": "Alpha", "dob": "01/01/1970"}, "beta": {"surname": "Beta", "dob": "02/02/1980"}}))
    (d / "text" / "a.txt").write_text("Patient Alpha seen 01/01/1970 DOB, follow up.")
    (d / "text" / "b.txt").write_text("CMS-1500 patient DOB 02021980 claim form.")   # names beta by DOB only, filed under alpha
    (d / "units" / "alpha.json").write_text(json.dumps([{"id": "a", "name": "note", "text_path": str(d / "text" / "a.txt")},
                                                       {"id": "b", "name": "claim", "text_path": str(d / "text" / "b.txt")}]))
    (d / "units" / "beta.json").write_text(json.dumps([]))
    found, checked, missing = cross_client.flags(d)
    assert checked == 2 and missing == [] and [f["file"] for f in found] == ["claim"] and found[0]["names"] == "beta"
    assert identity.run(sr) == 0
    assert any("IDENTITY? 'claim' routed to alpha but names beta" in line for line in log)
    assert cross_client.dob_variants("2/3/1990") == sorted(["02/03/1990", "02031990", "02-03-1990", "1990-02-03", "2/3/1990"])


def test_gate_modules_delegate_to_the_stages(job_dir: Path, firm_config_path: Path, data_root: Path) -> None:
    sr = _sr(job_dir, firm_config_path, data_root)
    d = sr.slug_dir
    (d / "units").mkdir(parents=True)
    (d / "units" / "alpha.json").write_text(json.dumps([{"id": "a", "name": "unexplained", "ext": ".pdf"}]))
    (d / "raw_manifest.jsonl").write_text(json.dumps({"id": "a", "name": "unexplained", "ext": ".pdf", "ok": True}) + "\n")
    (d / "runs" / "alpha").mkdir(parents=True)
    (d / "runs" / "alpha" / "entries_final.md").write_text("01/02/2026\nX | Medical Diagnoses\n\nY. (Exhibit 1 - p. 1)\n")
    (d / "out" / "alpha").mkdir(parents=True)
    (d / "out" / "alpha" / "page_map.json").write_text("[]")
    assert provenance.check(sr) == 1                                           # uncited, unexplained: holds
    (d / "runs" / "alpha" / "final-chronology.md").write_text("## Medical Chronology\n\nA claim long enough to audit here. (Exhibit 1 - p. 1)\n## Exhibit List\n")
    (d / "out" / "alpha" / "Exhibit 1 - X - 01-02-2026 (Medical Records).pdf").write_bytes(make_pdf([""]))
    ok, summary = claim_audit.check(d, "alpha")
    assert not ok and summary["never"] == 1                                    # never audited in final form
    (d / "out" / "alpha" / "page_map.json").unlink()                          # exhibits resolve by name without it
    (d / "nonrecord.json").write_text(json.dumps({"1": {"pages": 1, "blocks": [], "drop_pages": [1], "unknown": [], "cited_collision": [1]}}))
    assert extractive.dry_run(sr) == 1                                         # every cited page dropped
