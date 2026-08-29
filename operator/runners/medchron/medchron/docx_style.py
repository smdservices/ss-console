"""The document styling both deliverables share (the chronology and the
billing worksheet), so the two look like they came from the same firm.
Layout only; every character of content comes from the assembled markdown or
the chart JSON. The base font is the firm's (`format.font`)."""
from __future__ import annotations

from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLACK = RGBColor(0x1A, 0x1A, 0x1A)
CONTENT_W = 6.5


def _tcpr(cell: Any) -> Any:
    return cell._tc.get_or_add_tcPr()


def shade(cell: Any, fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    _tcpr(cell).append(el)


def borders(cell: Any, color: str = "424242", sz: int = 4) -> None:
    tcB = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        tcB.append(e)
    _tcpr(cell).append(tcB)


def vcenter(cell: Any) -> None:
    el = OxmlElement("w:vAlign")
    el.set(qn("w:val"), "center")
    _tcpr(cell).append(el)


def cell_margins(cell: Any, top: int = 60, bottom: int = 60, left: int = 90, right: int = 90) -> None:
    mar = OxmlElement("w:tcMar")
    for edge, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    _tcpr(cell).append(mar)


def fixed_layout(table: Any) -> None:
    el = OxmlElement("w:tblLayout")
    el.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(el)


def set_widths(table: Any, widths: list[float]) -> None:
    """Fixed layout honours w:gridCol, not per-cell w:tcW alone: set both and
    declare the table width, or the renderers fall back to even columns."""
    table.autofit = False
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tblW.set(qn("w:type"), "dxa")
    table._tbl.tblPr.append(tblW)
    for j, w in enumerate(widths[:len(table.columns)]):
        table.columns[j].width = Inches(w)
        for cell in table.columns[j].cells:
            cell.width = Inches(w)


def _cant_split(row: Any, val: str) -> None:
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), val)
    row._tr.get_or_add_trPr().append(el)


def allow_split(row: Any) -> None:
    _cant_split(row, "0")


def no_split(row: Any) -> None:
    _cant_split(row, "1")


def page_number_footer(doc: Any, font: str) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for tag, attrs, text in (("w:fldChar", {"w:fldCharType": "begin"}, None), ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
                             ("w:fldChar", {"w:fldCharType": "end"}, None)):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if text:
            el.text = text
        run._r.append(el)
    run.font.size = Pt(9)
    run.font.name = font
    run.font.color.rgb = BLACK


def build_styles(doc: Any, font: str) -> None:
    n = doc.styles["Normal"]
    n.font.name, n.font.size, n.font.color.rgb = font, Pt(12), BLACK
    n.paragraph_format.space_after = Pt(6)
    t = doc.styles["Title"]
    t.font.name, t.font.size, t.font.bold, t.font.color.rgb = font, Pt(18), True, BLACK
    if t.element.pPr is not None and t.element.pPr.find(qn("w:pBdr")) is not None:
        t.element.pPr.remove(t.element.pPr.find(qn("w:pBdr")))
    h = doc.styles["Heading 1"]
    h.font.name, h.font.size, h.font.bold, h.font.color.rgb = font, Pt(14), True, BLACK
    h.paragraph_format.space_before, h.paragraph_format.space_after = Pt(14), Pt(8)
    b = doc.styles["List Bullet"]
    b.font.name, b.font.size, b.font.color.rgb = font, Pt(10), BLACK
