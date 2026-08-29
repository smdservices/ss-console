"""`render`: the chronology .docx in the exemplar's layout. $0. Layout only:
every character of content comes from the assembled markdown.

The review note (Records Reviewed and Limitations) renders AT THE FRONT of
the client document, where the reader starts; the markdown keeps it at the
end because every stage before this one relies on that position. The
exhibit slice stops at that section, because prose after the exhibit table
once vanished from every delivered document (only pipe lines were kept).
The Prior Medical History block and the scope note are the body's preamble
and render before the entry table; markdown emphasis on their labels is
rendered as bold runs, not printed as asterisks. Timeline images render next
to the output document, never in a shared directory (one client's calendar
once rendered inside another's chronology).
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any

from .. import docx_style as S, timeline
from .base import StageRun
from .group import Canon

ENTRY_SPLIT = re.compile(r"(?m)^(?=\d{2}/\d{2}/\d{4}\s*(?:\(|$))")
DATE_HEAD = re.compile(r"^(\d{2})/(\d{2})/(\d{4})")
LIMITS = "## Records Reviewed and Limitations"


def emit_markdown_para(doc: Any, text: str) -> Any:
    p = doc.add_paragraph()
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text, flags=re.S)):
        if part:
            p.add_run(part).bold = bool(i % 2)
    return p


def write_cell(cell: Any, blocks: list[tuple[str, bool, bool]], font: str, size: int = 10) -> None:
    from docx.shared import Inches, Pt

    cell.text = ""
    first = cell.paragraphs[0]
    for i, (text, bold, bullet) in enumerate(blocks):
        p = first if i == 0 else cell.add_paragraph()
        if bullet:
            p.style = "List Bullet"
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.space_after = Pt(4)
        else:
            p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if text:
            r = p.add_run(text)
            r.bold, r.font.size, r.font.name, r.font.color.rgb = bool(bold), Pt(size), font, S.BLACK


def summary_table(doc: Any, rows: list[str], widths: list[float], font: str, header_fill: str = "f1f1f1",
                  bold_first_col: bool = True) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(re.fullmatch(r"-{2,}:?|:-{1,}:?", c or "---") for c in r)]
    if not cells:
        return
    ncols = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    S.fixed_layout(t)
    for i, row in enumerate(cells):
        S.no_split(t.rows[i])
        for j in range(ncols):
            c = t.cell(i, j)
            txt = row[j] if j < len(row) else ""
            if i > 0 and txt.startswith("(") and txt.endswith(")"):
                txt = txt[1:-1]
            write_cell(c, [(txt, i == 0 or (bold_first_col and j == 0), False)], font)
            S.borders(c, color="000000", sz=8)
            S.vcenter(c)
            S.cell_margins(c)
            if i == 0:
                S.shade(c, header_fill)
    S.set_widths(t, widths)
    doc.add_paragraph("")


def parse_entries(block: str) -> list[dict[str, Any]]:
    out = []
    for chunk in ENTRY_SPLIT.split(block):
        chunk = chunk.strip()
        m = DATE_HEAD.match(chunk)
        if not m:
            continue
        mm, dd, yy = m.groups()
        try:
            date(int(yy), int(mm), int(dd))
        except ValueError:
            continue
        lines = chunk.splitlines()
        out.append({"date": lines[0].strip(), "provider": lines[1].split("|")[0].strip() if len(lines) > 1 else "",
                    "body": lines[2:]})
    return out


def body_table(doc: Any, entries: list[dict[str, Any]], headings: set[str], font: str) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if not entries:
        return
    t = doc.add_table(rows=len(entries), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    S.fixed_layout(t)
    for i, e in enumerate(entries):
        left, right = t.cell(i, 0), t.cell(i, 1)
        write_cell(left, [(e["date"], True, False), (e["provider"], True, False)], font)
        S.shade(left, "f1f1f1")
        blocks: list[tuple[str, bool, bool]] = []
        prev_blank = True
        for ln in e["body"]:
            s = ln.strip()
            if not s:
                if not prev_blank:
                    blocks.append(("", False, False))
                    prev_blank = True
                continue
            is_head = s in headings
            blocks.append((s, is_head, not is_head))
            prev_blank = False
        while blocks and blocks[-1][0] == "":
            blocks.pop()
        write_cell(right, blocks or [("", False, False)], font)
        for c in (left, right):
            S.borders(c, color="424242", sz=4)
            S.vcenter(c)
            S.cell_margins(c)
        S.allow_split(t.rows[i])
    S.set_widths(t, [1.38, 5.09])


def gap_bar(doc: Any, a: date, b: date, days: int, font: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f"Treatment Gap  {a.strftime('%b %-d, %Y')} - {b.strftime('%b %-d, %Y')}  ({days} days)")
    r.bold, r.font.size, r.font.name, r.font.color.rgb = True, Pt(9), font, S.BLACK
    bdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), "9a9a9a")
        bdr.append(e)
    p._p.get_or_add_pPr().append(bdr)


def section(text: str, name: str, upto: str) -> str:
    return text.split(f"## {name}\n", 1)[1].split(f"## {upto}", 1)[0]


def run(sr: StageRun) -> int:
    from docx import Document
    from docx.shared import Inches

    rd = sr.slug_dir / "runs" / sr.unit.unit
    out_dir = sr.slug_dir / "out" / sr.unit.unit
    text = (rd / "final-chronology.md").read_text(encoding="utf-8")
    font = str(sr.cfg.get("format", "font") or "Calibri")
    headings = {str(h) for h in (sr.cfg.get("format", "subsections") or [])} | {"Discharge / Discontinuation in Care"}
    title = text.splitlines()[0].strip()
    prov_rows = [ln for ln in section(text, "Treatment Timeline", "Diagnostic Highlights").splitlines() if ln.strip().startswith("|")]
    icd_rows = [ln for ln in section(text, "Diagnostic Highlights", "Medical Chronology").splitlines() if ln.strip().startswith("|")]
    body = section(text, "Medical Chronology", "Exhibit List")
    ex_block = text.split("## Exhibit List", 1)[1].split(LIMITS, 1)[0]
    ex_rows = [ln for ln in ex_block.splitlines() if ln.strip().startswith("|")]
    limits_block = text.split(LIMITS, 1)[1] if LIMITS in text else ""
    incident = date(*(int(x) for x in sr.job.incident_date.split("-")))
    imgdir = out_dir / "img"
    visuals = timeline.render(text, imgdir, incident, int(sr.cfg.get("chronology", "treatment_gap_days", 45)),
                              Canon(sr.cfg), {h.lower() for h in headings})

    doc = Document()
    S.build_styles(doc, font)
    S.page_number_footer(doc, font)
    doc.add_paragraph(title, style="Title")
    limits_emitted = 0
    if limits_block.strip():
        doc.add_paragraph("Records Reviewed and Limitations", style="Heading 1")
        for raw_line in limits_block.splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#"):
                continue
            if line.startswith("* "):
                emit_markdown_para(doc, line[2:].strip()).style = doc.styles["List Bullet"]
            else:
                emit_markdown_para(doc, line)
            limits_emitted += 1
        if not limits_emitted:
            sr.log("build_doc emitted a limitations section and the renderer produced no paragraphs from it")
            return 1
        doc.add_paragraph("")
    doc.add_paragraph("1.  Treatment Timeline", style="Heading 1")
    if visuals["lane_order"]:
        doc.add_picture(str(imgdir / "timeline.png"), width=Inches(S.CONTENT_W))
        for y in visuals["years"]:
            doc.add_picture(str(imgdir / f"cal-{y}.png"), width=Inches(S.CONTENT_W))
            for a, b, n in visuals["gaps"]:
                if a.year == y:
                    gap_bar(doc, a, b, n, font)
        doc.add_picture(str(imgdir / "legend.png"), width=Inches(S.CONTENT_W))
    doc.add_paragraph("2.  Past Treatment History", style="Heading 1")
    summary_table(doc, prov_rows, [2.54, 1.74, 1.10, 1.12], font)
    doc.add_paragraph("3.  Diagnostic Highlights", style="Heading 1")
    summary_table(doc, icd_rows, [0.81, 3.43, 1.12, 1.12], font)
    doc.add_paragraph("4.  Medical Summary", style="Heading 1")
    preamble = ENTRY_SPLIT.split(body)[0].strip()
    if preamble:
        for para in [p.strip() for p in preamble.split("\n\n") if p.strip()]:
            first = para.splitlines()[0].strip()
            if first and len(first) < 60 and not first.endswith("."):
                doc.add_paragraph(first, style="Heading 2")
                rest = "\n".join(para.splitlines()[1:]).strip()
                if rest:
                    emit_markdown_para(doc, rest)
            else:
                emit_markdown_para(doc, para)
        doc.add_paragraph("")
    entries = parse_entries(body)
    body_table(doc, entries, headings, font)
    doc.add_paragraph("")
    doc.add_paragraph("5.  Exhibit List", style="Heading 1")
    summary_table(doc, ex_rows, [1.25, 5.25], font, header_fill="f3f6fe", bold_first_col=False)
    stamp = sr.date_stamp or date.today().strftime("%m-%d-%y")
    name = str(sr.cfg.get("delivery", "chronology_name_template") or "{CLIENT} - Medical Chronology {MM-DD-YY}.docx")
    dst = out_dir / name.replace("{CLIENT}", sr.unit.client_name).replace("{MM-DD-YY}", stamp)
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    sr.log(f"wrote {dst.name}: {len(entries)} entries, {max(len(prov_rows) - 1, 0)} providers, {max(len(icd_rows) - 1, 0)} ICD rows, "
           f"{max(len(ex_rows) - 1, 0)} exhibits, {len(visuals['years'])} calendar year(s), {len(visuals['gaps'])} gap bar(s), "
           f"{limits_emitted} limitations line(s)")
    return 0
