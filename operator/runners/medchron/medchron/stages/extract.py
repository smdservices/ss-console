"""`extract`: local text extraction over the pulled `raw/` set. $0.

Per file: pypdf text with `[p.N]` page markers, docx via python-docx. A PDF
whose text layer is absent, is glyph indices, or is a constant-offset cipher
goes to the scan queue for the vision pass; the two detectors and their
thresholds are calibrated on delivered runs (see the docstrings), and both
failure shapes produced a chronology that looked complete while omitting or
understating a record, which is worse than omitting the file.

Output (the shapes every later stage reads):
  extracted.jsonl   one record per file: {id, name, folder, ext, pages, chars,
                    text_path | scan: true}
  text/<id>.txt     the text with [p.N] markers
  scan_queue.json   files needing the vision pass
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from .base import StageRun, read_jsonl

GLYPH_RUN = re.compile(r"(?:/\d+ ?){4,}")
GLYPH_THRESHOLD = 0.08     # junk files score 0.21-0.31; the highest clean file 0.031
ENGLISH_FLOOR = 0.02       # ciphered files 0.0000; lowest legitimate (a bill) 0.0317
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tif", ".tiff"}
WORD_EXTS = {".docx", ".doc"}

# A short, deliberately boring list: words that appear in ANY page of English
# prose or clinical narrative, including tables and forms.
_STOP = frozenset(
    "the and of to in for with was were is are on at by a an patient this that "
    "no not his her he she had has been from as or be date name page report".split())


def pdf_text(path: Path) -> list[tuple[int, str]]:
    from pypdf import PdfReader

    pages = []
    for i, pg in enumerate(PdfReader(str(path)).pages, 1):
        try:
            t = pg.extract_text() or ""
        except Exception:  # noqa: BLE001 - one unreadable page is an empty page
            t = ""
        pages.append((i, t))
    return pages


def docx_text(path: Path) -> str:
    import docx

    d = docx.Document(str(path))
    out = [p.text for p in d.paragraphs if p.text.strip()]
    for tb in d.tables:
        for row in tb.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(out)


def not_english(txt: str) -> bool:
    """A text layer that decodes to something that is not language: a
    constant character offset ("6HH$GGHQGXP%HORZ" for "See Addendum Below")
    contains no slashes and passes every structural check, and its digits
    decode to punctuation, so a composed entry reports the study as "not
    legible" while the page states it plainly."""
    toks = re.findall(r"[A-Za-z]{2,}", txt.lower())
    if len(toks) < 50:
        return False
    return sum(1 for w in toks if w in _STOP) / len(toks) < ENGLISH_FLOOR


def glyph_junk(txt: str) -> bool:
    """A text layer of glyph INDICES ("/0/1/2/3"): a subsetted font with no
    ToUnicode map. Passes every has-a-text-layer test; unreadable to the model."""
    if not txt:
        return False
    covered = sum(len(m.group()) for m in GLYPH_RUN.finditer(txt))
    return covered / len(txt) > GLYPH_THRESHOLD


def _text_layer_absent(total: int, pages: int) -> bool:
    return total < 40 * max(1, pages) * 0.05 and total < 400


def live_rows(rows: list[dict[str, Any]]) -> tuple[dict[str, dict[str, Any]], int]:
    """Last write wins (retries append); a byte-duplicate row retires its id."""
    ok: dict[str, dict[str, Any]] = {}
    dupes = 0
    for r in rows:
        if r.get("ok"):
            if r.get("duplicate_of"):
                dupes += 1
                ok.pop(r["id"], None)
                continue
            ok[r["id"]] = r
    return ok, dupes


def _extract_one(d: Path, r: dict[str, Any], scans: list[dict[str, Any]]) -> dict[str, Any]:
    ext = (r.get("ext") or "").lower()
    path = Path(r["path"])
    rec: dict[str, Any] = {"id": r["id"], "name": r["name"], "folder": r["folder"], "ext": ext}
    text_path = d / "text" / f"{r['id']}.txt"
    try:
        if ext == ".pdf":
            pages = pdf_text(path)
            total = sum(len(t.strip()) for _, t in pages)
            rec["pages"] = len(pages)
            txt = "\n".join(f"[p.{n}]\n{t}" for n, t in pages)
            junk = glyph_junk(txt) or not_english(txt)
            if _text_layer_absent(total, len(pages)) or junk:
                rec["scan"] = True
                if junk:
                    rec["glyph_junk"] = True
                    # A stale text file from a previous extract satisfies the
                    # vision pass's resume check, so routing to vision means
                    # removing the text that routed it there.
                    text_path.unlink(missing_ok=True)
                scans.append(rec)
            else:
                text_path.write_text(txt, encoding="utf-8")
                rec["chars"] = len(txt)
                rec["text_path"] = str(text_path)
                empty = [n for n, t in pages if not t.strip()]
                if empty:
                    rec["empty_pages"] = empty
        elif ext in WORD_EXTS:
            try:
                txt = docx_text(path)
            except Exception as exc:  # noqa: BLE001
                rec["error"] = f"docx: {exc}"[:120]
                txt = ""
            if txt:
                text_path.write_text(txt, encoding="utf-8")
                rec["chars"] = len(txt)
                rec["text_path"] = str(text_path)
            else:
                rec.setdefault("error", "no text")
        elif ext in IMAGE_EXTS:
            rec["scan"] = True
            rec["image"] = True
            scans.append(rec)
        else:
            rec["error"] = "unhandled ext"
    except Exception as exc:  # noqa: BLE001 - one broken file is one row
        rec["error"] = str(exc)[:200]
    return rec


def run(sr: StageRun) -> int:
    d = sr.slug_dir
    (d / "text").mkdir(parents=True, exist_ok=True)
    ok, dupes = live_rows(read_jsonl(d / "raw_manifest.jsonl"))
    scans: list[dict[str, Any]] = []
    with (d / "extracted.jsonl").open("w", encoding="utf-8") as out:
        for i, r in enumerate(ok.values(), 1):
            out.write(json.dumps(_extract_one(d, r, scans)) + "\n")
            if i % 40 == 0:
                sr.log(f"{i}/{len(ok)}")
    (d / "scan_queue.json").write_text(json.dumps(scans, indent=1), encoding="utf-8")
    sr.log(f"DONE {len(ok)} files, {len(scans)} to scan queue, {dupes} byte-duplicates skipped")
    return 0
