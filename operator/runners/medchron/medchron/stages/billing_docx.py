"""`billing_docx`: the billing chart rendered as a worksheet matching the
chronology. $0.

Content rules, same as the chronology's: every figure names the document it
came from; nothing is computed that is not shown; a provider whose total is
not derivable says so instead of carrying a plausible number; the grand
total is labelled SUBTOTAL, with a warning, whenever any row is missing. A
worksheet is where a misattributed figure does damage, because it reads as
the specials number: it refuses (exit 1) a chart built without the patient
filter on a joint matter, or one carrying a suspect (lost-decimal) amount.
"""
from __future__ import annotations

from datetime import date
from typing import Any

from .. import docx_style as S
from .base import StageRun, read_json


def money(v: float | None) -> str:
    return f"${v:,.2f}" if v is not None else "not derivable"


def dshow(d: str | None) -> str:
    return f"{d[5:7]}/{d[8:10]}/{d[0:4]}" if d else "-"


def _cell_text(c: Any, text: str, size: float = 9.5, bold: bool = False, right: bool = False) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = c.paragraphs[0]
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
    if right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(text)
    r.font.size, r.bold, r.font.color.rgb = Pt(size), bold, S.BLACK


def table(doc: Any, header: list[str], rows: list[list[str]], widths: list[float], total_row: list[str] | None = None,
          right_cols: tuple[int, ...] = ()) -> None:
    t = doc.add_table(rows=len(rows) + 1 + (1 if total_row else 0), cols=len(header))
    S.fixed_layout(t)
    S.set_widths(t, widths)

    def fill(cells: Any, values: list[str], bold: bool, shaded: bool) -> None:
        for j, v in enumerate(values):
            c = cells[j]
            if shaded:
                S.shade(c, "f1f1f1")
            S.borders(c)
            S.cell_margins(c)
            S.vcenter(c)
            _cell_text(c, str(v), bold=bold, right=j in right_cols)

    fill(t.rows[0].cells, header, True, True)
    for i, row in enumerate(rows, 1):
        S.no_split(t.rows[i])
        fill(t.rows[i].cells, row, False, False)
    if total_row:
        S.no_split(t.rows[-1])
        fill(t.rows[-1].cells, total_row, True, True)


def para(doc: Any, text: str, size: float = 9.5, italic: bool = False, before: int = 8) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(before), Pt(2)
    r = p.add_run(text)
    r.font.size, r.italic, r.font.color.rgb = Pt(size), italic, S.BLACK


def run(sr: StageRun) -> int:
    from docx import Document
    from docx.shared import Inches

    d = sr.slug_dir
    unit = sr.unit.unit
    sfx = f"-{unit}" if sr.job.joint else ""
    data = read_json(d / f"billing_chart{sfx}.json", None)
    if data is None:
        sr.log("billing_chart.json not found; billing_chart runs first")
        return 1
    if sr.job.joint and not data.get("patient_filter"):
        sr.log("chart was built without a patient filter on a joint matter; refusing to render a worksheet")
        return 1
    if data.get("suspect_amounts"):
        sr.log(f"chart carries {len(data['suspect_amounts'])} suspect amount(s) (likely lost decimal); refusing to render")
        return 1
    rows = data["rows"]
    if not rows and not data.get("subrogation") and not data.get("vendor_invoices"):
        sr.log("no billing rows and no excluded items; no worksheet for this unit")
        return 0
    gaps = [r for r in rows if r["total"] is None]
    font = str(sr.cfg.get("format", "font") or "Calibri")
    doc = Document()
    S.build_styles(doc, font)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = s.bottom_margin = Inches(0.8)
    S.page_number_footer(doc, font)
    who = sr.unit.client_name
    doc.add_paragraph(f"{who} - Medical Billing Summary", style="Title")
    para(doc, f"Prepared {date.today():%B %-d, %Y}. Providers and treatment dates are taken from the medical chronology "
              f"for this matter. Each amount names the document it was read from. No figure here was calculated by us "
              f"except the total, which is the sum of the rows shown.", italic=True)
    if gaps:
        para(doc, f"INCOMPLETE: {len(gaps)} of {len(rows)} providers have no total that can be supported from the documents "
                  f"on file. The subtotal below is therefore NOT this case's medical specials figure. See \"What Is Needed\" "
                  f"at the end.")
    doc.add_paragraph("Billing by Provider", style="Heading 1")
    body: list[list[str]] = []
    grand, orphan = 0.0, False
    for r in rows:
        if r["total"] is not None:
            grand += r["total"]
        name = r["provider"]
        if r["first"] is None and r["last"] is None and r.get("lien"):
            name, orphan = name + " *", True
        body.append([name, dshow(r["first"]), dshow(r["last"]), r["basis"], money(r["total"])])
    table(doc, ["Provider", "First DOS", "Last DOS", "Source", "Total Billed"], body, [2.3, 0.8, 0.8, 2.35, 1.05],
          total_row=["SUBTOTAL" if gaps else "GRAND TOTAL", "", "", "", money(round(grand, 2))], right_cols=(4,))
    if orphan:
        para(doc, "* This line appears on the firm's lien report with no provider name printed, and no matching provider or "
                  "treatment appears in the medical chronology. It is included so this summary reconciles to the firm's "
                  "own report, but it cannot be attributed to a provider from the documents on file.", size=9)
    alts = [r for r in rows if r.get("alternates")]
    if alts:
        doc.add_paragraph("Where Documents Disagree", style="Heading 1")
        para(doc, "More than one figure appears in the file for these providers. The chart uses the one named in the Source "
                  "column; the others are listed so the choice is visible.")
        table(doc, ["Provider", "Figures found in the file"],
              [[r["provider"], ", ".join(f"${a:,.2f}" for a in r["alternates"])] for r in alts], [2.6, 4.6])
    if data.get("subrogation"):
        doc.add_paragraph("Health-Plan and Subrogation Claims", style="Heading 1")
        para(doc, "These are claims by a health plan (or its recovery vendor) for benefits it paid, not charges by a treating "
                  "provider. They are NOT included in the total above, and including them would double-count the underlying care.")
        seen: set = set()
        srows = []
        for fl, pv, amts in data["subrogation"]:
            k = (str(pv)[:40], fl[:40])
            if k not in seen:
                seen.add(k)
                srows.append([str(pv)[:40], fl[:40], ", ".join(str(a) for a in (amts or [])[:2])])
        table(doc, ["Claimant", "Source Document", "Amounts Printed"], srows[:14], [2.2, 3.0, 2.0])
    if data.get("vendor_invoices"):
        doc.add_paragraph("Excluded: Vendor Invoices", style="Heading 1")
        para(doc, "Invoices billed to the firm by service vendors (record retrieval and similar). Not medical specials.")
        seen2: set = set()
        vrows = []
        for fl, pv in data["vendor_invoices"]:
            if fl[:40] not in seen2:
                seen2.add(fl[:40])
                vrows.append([str(pv)[:36], fl[:44]])
        table(doc, ["Vendor", "Source Document"], vrows[:14], [2.6, 4.6])
    if data.get("quarantined"):
        doc.add_paragraph("Excluded: Documents Naming Another Patient", style="Heading 1")
        para(doc, "This matter holds documents for more than one person. The following name a different patient and are "
                  "excluded from every figure above.")
        seen3: set = set()
        qrows = []
        for fl, wh, pv in data["quarantined"]:
            k = (str(wh)[:26], fl[:34])
            if k not in seen3:
                seen3.add(k)
                qrows.append([str(wh)[:26], str(pv)[:28], fl[:36]])
        table(doc, ["Patient Named", "Provider", "Source Document"], qrows[:16], [1.8, 2.2, 3.2])
    if gaps:
        doc.add_paragraph("What Is Needed", style="Heading 1")
        para(doc, "For the providers below, the file holds individual claim forms but no account ledger or statement of the "
                  "total billed. Summing the claim forms is not reliable here because the same bills appear both individually "
                  "and inside compiled record sets, so a sum would count them twice. A current ledger from each provider, "
                  "or a Lienholder and Balances report for this matter, would complete the chart.")
        table(doc, ["Provider", "What is on file"], [[r["provider"], r["basis"]] for r in gaps], [2.6, 4.6])
    if data.get("failed_pages"):
        doc.add_paragraph("Pages Not Read", style="Heading 1")
        para(doc, f"{len(data['failed_pages'])} page(s) could not be read. Figures above are a floor until they are resolved.")
    stamp = sr.date_stamp or date.today().strftime("%m-%d-%y")
    glob = str(sr.cfg.get("delivery", "worksheet_glob") or "* - Medical Billing Worksheet *.docx")
    out = d / "out" / unit / glob.replace("*", who, 1).replace("*", stamp, 1)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    sr.log(f"wrote {out.name}: {len(rows)} provider(s), {'subtotal' if gaps else 'total'} {money(round(grand, 2))}"
           + (f", {len(gaps)} gap(s)" if gaps else ""))
    return 0
