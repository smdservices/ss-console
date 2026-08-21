"""Render drafting content INTO a firm's Word template: code owns typography.

This is the typographic half of ADR 0083 ("the model fills content into a
structure; code owns typography and required elements"). ``docgrammar`` gives
us blocks; this module decides how each block LOOKS for a given document class,
and it takes that decision from exactly one place when it can: **the firm's own
Word template**, a ``.docx`` in the firm's Document Library. The renderer opens
that file as the base document (page setup, headers/footers/letterhead, styles,
numbering and theme all survive), clears its body, and writes the content back
in using a small contract of NAMED PARAGRAPH STYLES. A template that defines a
named style wins; a template that lacks one gets the class's product default
applied inline, and the fallback is reported. Nothing is written back into the
firm's file, ever.

Why typography lives ONLY in the .docx: a firm edits a style in Word and the
next draft honors it. No config publish, no reboot, no SMD in the loop. That is
the difference between "formatted by us" and "formatted the way the firm
formats", and it is the same sentence as "the firm authors its own posture".

Why this module never invents legal content: item numbers come from the
propounded set; the 35-interrogatory rule is an aggregate across the matter and
attorney-reserved; a statute-bound declaration is jurisdiction-specific. The
model writes labels, numerals, captions, signature blocks, and proof of
service exactly as the skeletons do today; this module STYLES them (a line that
looks like an item label gets the label style) and adds nothing. Content stays
under the drafting gates; typography is code.

Named style contract (a firm's template may define any subset):
  SMD Body, SMD Item Label, SMD Item Text, SMD Heading 1/2/3, SMD Caption,
  SMD Signature

Refusals (the only ones): a multi-section base document (the body clear would
silently keep the LAST section's page setup and drop the rest; a wrong
letterhead on a client letter is worse than an honest "not yet").
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from . import docgrammar as g
from .docx_base import has_style, open_as_base, usable_paragraph_style
from .docx_format_types import (
    DEFAULT_FONT,
    DEFAULT_SIZE_PT,
    NAMED_STYLES,
    ROLE_FALLBACK,
    FormatRefused,
    FormatReport,
)

# ---- Document classes and their styling rules --------------------------------

DOCUMENT_CLASSES = (
    "discovery_set",
    "discovery_response",
    "demand_letter",
    "mediation_brief",
    "memo",
    "letter",
)

# Label shapes the renderer STYLES (it never writes them). Anchored at line
# start, whole-line by construction of the skeletons ("**SPECIAL INTERROGATORY
# NO. 7:**"). A mid-paragraph mention is prose and stays body text.
_DISCOVERY_LABEL = re.compile(
    r"^(?:RESPONSE TO |SUPPLEMENTAL RESPONSE TO )?"
    r"(?:SPECIAL INTERROGATOR(?:Y|IES)|FORM INTERROGATOR(?:Y|IES)|"
    r"REQUESTS? FOR (?:PRODUCTION|ADMISSION|INSPECTION)|INSPECTION DEMANDS?|"
    r"DEMANDS? FOR (?:PRODUCTION|INSPECTION))"
    r"\s+NO\.?\s*\S",
    re.IGNORECASE,
)
_DEFINITION_LABEL = re.compile(r"^DEFINITIONS?\b", re.IGNORECASE)
_LABEL_MAX_CHARS = 90


@dataclass(frozen=True)
class ClassRules:
    """Per-class styling decisions. Typography values here are the PRODUCT
    DEFAULTS used only when the base template lacks the named style; they are
    deliberately the authored standards of the first engagements (plain Word,
    not pleading paper) and are labeled as starters the firm edits in Word."""

    label_patterns: tuple[re.Pattern[str], ...] = ()
    # Paragraphs following a label (until the next label/heading/table) are item
    # text: first-line indented, with the "between items" spacing as spacing
    # after the paragraph (the authored standard: double-spaced BETWEEN requests,
    # not more).
    item_text: bool = False
    caption_table_first: bool = False  # first table is the caption (court docs)
    heading_align: tuple[str, str, str] = ("center", "left", "left")
    heading_underline: tuple[bool, bool, bool] = (False, True, False)
    heading_indent_in: tuple[float, float, float] = (0.0, 0.5, 1.0)
    page_numbers: bool = True
    body_line_spacing: float = 1.0  # multiple
    item_line_spacing: float = 2.0
    item_space_after_pt: float = 0.0
    first_line_indent_in: float = 0.5


CLASS_RULES: dict[str, ClassRules] = {
    "discovery_set": ClassRules(
        label_patterns=(_DISCOVERY_LABEL, _DEFINITION_LABEL),
        item_text=True,
        caption_table_first=True,
        item_line_spacing=2.0,
        item_space_after_pt=12.0,
    ),
    "discovery_response": ClassRules(
        label_patterns=(_DISCOVERY_LABEL, _DEFINITION_LABEL),
        item_text=True,
        caption_table_first=True,
        item_line_spacing=2.0,
        item_space_after_pt=12.0,
    ),
    "demand_letter": ClassRules(page_numbers=True, heading_align=("left", "left", "left")),
    "mediation_brief": ClassRules(
        caption_table_first=True,
        heading_align=("center", "left", "left"),
        heading_underline=(False, True, False),
        body_line_spacing=2.0,
    ),
    "memo": ClassRules(heading_align=("left", "left", "left"), page_numbers=True),
    "letter": ClassRules(heading_align=("left", "left", "left"), page_numbers=False),
}

# ---- Rendering -----------------------------------------------------------------


def render_document(
    markdown: str,
    document_class: str,
    base_bytes: bytes | None,
    report: FormatReport | None = None,
) -> tuple[bytes, FormatReport]:
    """Render ``markdown`` for ``document_class`` into ``base_bytes`` (the firm's
    template) or the stock starter base. Pure: no network, no client."""
    if document_class not in CLASS_RULES:
        raise ValueError(f"unknown document_class {document_class!r}; known: {', '.join(DOCUMENT_CLASSES)}")
    rules = CLASS_RULES[document_class]
    report = report or FormatReport(document_class=document_class)
    doc = open_as_base(base_bytes, report)
    writer = _Writer(doc, rules, report)
    for block in g.parse_document(markdown):
        writer.write(block)
    if rules.page_numbers:
        writer.ensure_page_number()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), report


class _Writer:
    def __init__(self, doc, rules: ClassRules, report: FormatReport) -> None:
        self.doc = doc
        self.rules = rules
        self.report = report
        self.in_item = False
        self.tables_seen = 0

    # -- style application ------------------------------------------------------

    def _resolve_style(self, role: str) -> tuple[str | None, bool]:
        """``(style to apply, delegated?)`` for a role, or ``(None, False)`` to
        format inline.

        Three states, in order. Our named style if the base defines it. Else the
        base's OWN conventional equivalent, when the role has one — that is what
        makes "edit the style in Word and the next draft follows" true for a
        firm template that never heard of our names. Else inline.
        """
        if usable_paragraph_style(self.doc, role):
            self.report.styles_honored.append(role)
            return role, False
        alt = ROLE_FALLBACK.get(role)
        if alt and usable_paragraph_style(self.doc, alt):
            self.report.styles_delegated[role] = alt
            return alt, True
        self.report.fallbacks.append(role)
        if alt:
            # An instruction the firm can act on, not a diagnostic string.
            note = f"the base defines neither {role!r} nor {alt!r}; add {alt!r} in Word to control this level"
            if note not in self.report.notes:
                self.report.notes.append(note)
        return None, False

    def _para(self, style_name: str):
        """Add a paragraph for ``style_name``.

        Returns ``(paragraph, styled)`` where ``styled`` is True only when our
        OWN named style carried it. A delegated paragraph reports ``False`` on
        purpose: the firm's style supplies font and colour, but the class's
        layout (indents, the double-spacing between discovery items, a centered
        pleading heading) is a court requirement code still owns, so the caller
        must keep applying it. See ``_heading``.
        """
        applied, delegated = self._resolve_style(style_name)
        if applied is None:
            return self.doc.add_paragraph(), False
        try:
            return self.doc.add_paragraph(style=applied), not delegated
        except (KeyError, ValueError):
            # A style that resolved but will not apply degrades to inline; it
            # never kills the render.
            self.report.fallbacks.append(style_name)
            self.report.styles_delegated.pop(style_name, None)
            return self.doc.add_paragraph(), False

    def _runs(self, para, runs: tuple[g.Run, ...], *, caps: bool = False, bold: bool = False, underline: bool = False) -> None:
        for r in runs:
            run = para.add_run(r.text)
            if r.marker:
                continue  # literal, unstyled, render-visible
            if r.bold or bold:
                run.bold = True
            if r.italic:
                run.italic = True
            if underline:
                run.underline = True
            if caps:
                run.font.all_caps = True

    # -- blocks ------------------------------------------------------------------

    def write(self, block: g.Block) -> None:
        if isinstance(block, g.Heading):
            self._heading(block)
        elif isinstance(block, g.Paragraph):
            self._paragraph(block)
        elif isinstance(block, g.Bullet):
            self._bullet(block)
        elif isinstance(block, g.Numbered):
            self._numbered(block)
        elif isinstance(block, g.Table):
            self._table(block)
        elif isinstance(block, g.HRule):
            self._hrule()

    def _heading(self, block: g.Heading) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt

        self.in_item = False
        self.report.blocks_styled["headings"] += 1
        level = block.level
        role = f"SMD Heading {level}"
        para, styled = self._para(role)
        if styled:
            self._runs(para, block.runs)
            return
        delegated = role in self.report.styles_delegated
        # Layout stays ours even when the firm's own heading style carries the
        # typography: a mediation brief's level-1 heading is CENTERED because
        # the court expects it there, and a firm's built-in Heading 1 is
        # left-aligned. Font, size and colour come from their style; placement
        # does not.
        align = self.rules.heading_align[level - 1]
        para.alignment = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
        para.paragraph_format.left_indent = Inches(self.rules.heading_indent_in[level - 1])
        para.paragraph_format.keep_with_next = True
        if delegated:
            # Emphasis is the firm style's business; forcing bold/underline over
            # it would defeat the edit we just made possible.
            self._runs(para, block.runs)
            return
        para.paragraph_format.space_before = Pt(12 if level == 1 else 6)
        self._runs(para, block.runs, bold=True, underline=self.rules.heading_underline[level - 1])

    def _is_label(self, text: str) -> bool:
        # A label is a SHORT line that starts with a label phrase ("SPECIAL
        # INTERROGATORY NO. 7:"); a prose sentence that happens to open with the
        # same words is body text. The length cap is the tie-breaker.
        t = text.strip().strip("*").strip()
        return len(t) <= _LABEL_MAX_CHARS and any(p.match(t) for p in self.rules.label_patterns)

    def _paragraph(self, block: g.Paragraph) -> None:
        from docx.shared import Inches, Pt

        if self._is_label(block.text):
            self.report.blocks_styled["labels"] += 1
            self.in_item = True
            para, styled = self._para("SMD Item Label")
            if styled:
                self._runs(para, block.runs)
            else:
                para.paragraph_format.space_before = Pt(12)
                self._runs(para, block.runs, caps=True, bold=True, underline=True)
            return
        if self.in_item and self.rules.item_text:
            para, styled = self._para("SMD Item Text")
            if not styled:
                pf = para.paragraph_format
                pf.first_line_indent = Inches(self.rules.first_line_indent_in)
                pf.line_spacing = self.rules.item_line_spacing
                pf.space_after = Pt(self.rules.item_space_after_pt)
            self._runs(para, block.runs)
            return
        para, styled = self._para("SMD Body")
        if not styled and self.rules.body_line_spacing != 1.0:
            para.paragraph_format.line_spacing = self.rules.body_line_spacing
        self._runs(para, block.runs)

    def _bullet(self, block: g.Bullet) -> None:
        from docx.shared import Inches

        if usable_paragraph_style(self.doc, "List Bullet"):
            try:
                para = self.doc.add_paragraph(style="List Bullet")
                self.report.styles_honored.append("List Bullet")
                self._runs(para, block.runs)
                return
            except (KeyError, ValueError):
                self.report.fallbacks.append("List Bullet")
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.add_run("•\t")
        self._runs(para, block.runs)

    def _numbered(self, block: g.Numbered) -> None:
        from docx.shared import Inches

        para, _ = self._para("SMD Body")
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.add_run(f"{block.label}\t")
        self._runs(para, block.runs)

    def _table(self, block: g.Table) -> None:
        self.in_item = False
        self.tables_seen += 1
        self.report.blocks_styled["tables"] += 1
        ncols = max(len(row) for row in block.rows)
        table = self.doc.add_table(rows=len(block.rows), cols=ncols)
        caption = self.rules.caption_table_first and self.tables_seen == 1
        _set_table_borders(table, inside_vertical_only=caption)
        for r_idx, row in enumerate(block.rows):
            for c_idx in range(ncols):
                cell = table.cell(r_idx, c_idx)
                runs = row[c_idx] if c_idx < len(row) else ()
                para = cell.paragraphs[0]
                if caption and usable_paragraph_style(self.doc, "SMD Caption"):
                    try:
                        para.style = self.doc.styles["SMD Caption"]
                        self.report.styles_honored.append("SMD Caption")
                    except (KeyError, ValueError):
                        self.report.fallbacks.append("SMD Caption")
                self._runs(para, runs, bold=(block.header and r_idx == 0))
        # Breathing room after a table: an empty body paragraph.
        self.doc.add_paragraph()

    def _hrule(self) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        self.in_item = False
        para = self.doc.add_paragraph()
        ppr = para._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pbdr.append(bottom)
        ppr.append(pbdr)

    # -- footer --------------------------------------------------------------------

    def ensure_page_number(self) -> None:
        """Add a centered PAGE field to the footer when the base has no footer
        text at all. A base with its own footer (a letterhead's address line, a
        firm's own page numbering) is left exactly as the firm built it."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        section = self.doc.sections[0]
        footer = section.footer
        existing = "".join(p.text for p in footer.paragraphs).strip()
        if existing or _footer_has_field(footer):
            return
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
            if kind:
                el = OxmlElement("w:fldChar")
                el.set(qn("w:fldCharType"), kind)
            else:
                el = OxmlElement("w:instrText")
                el.set(qn("xml:space"), "preserve")
                el.text = f" {text} "
            run._r.append(el)
        self.report.notes.append("page number field added to the footer")


def _footer_has_field(footer) -> bool:
    from docx.oxml.ns import qn

    return footer._element.find(f".//{qn('w:fldChar')}") is not None or footer._element.find(f".//{qn('w:fldSimple')}") is not None


def _set_table_borders(table, *, inside_vertical_only: bool) -> None:
    """Explicit ``w:tblBorders`` via lxml: a firm template rarely defines
    ``Table Grid``, so never rely on a table style existing. A caption table
    gets the classic look (a vertical rule between the columns, nothing else);
    every other table gets a thin grid."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    edges = ("insideV",) if inside_vertical_only else ("top", "left", "bottom", "right", "insideH", "insideV")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    for old in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(old)
    tbl_pr.append(borders)


__all__ = [
    "CLASS_RULES",
    "DEFAULT_FONT",
    "DEFAULT_SIZE_PT",
    "DOCUMENT_CLASSES",
    "NAMED_STYLES",
    "ROLE_FALLBACK",
    "ClassRules",
    "FormatRefused",
    "FormatReport",
    "has_style",
    "open_as_base",
    "render_document",
    "usable_paragraph_style",
]
