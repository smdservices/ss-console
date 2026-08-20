"""Text extraction for matter documents fetched via ``read_document``.

Bounded, dependency-light, and deliberately dumb: the goal is to hand the
agent the document's TEXT as data — never to render, execute, or interpret
anything inside it. Supported types are the three that cover a law office's
matter folder (PDF, DOCX, plain text); everything else returns an explicit
unsupported error so the skill can fail closed instead of guessing.

SCANNED PAPER (ss#2464). A law firm's record is half photocopies: on the
Robertus matter, 24 of 104 PDFs carry no text layer at all, and every one of
them is clinical — MRI reports, office notes, the X-ray reads. pypdf returns
``""`` for those, which used to leave ``read_document`` reporting
``total_chars: 0`` with no error: a skill could not tell an empty document
from an unreadable one. Two things changed here, and the second is the whole
discipline:

1. ``extract_text_ex`` returns an :class:`ExtractResult` that NAMES the path
   the text came from (or the reason there is none). A marker is never text —
   ``ExtractResult.text`` is empty for every non-success method, and ``reason``
   is drawn from a closed set, never from an exception string or model prose.
2. A scanned PDF can be TRANSCRIBED by a vision read (``vision.py``), but only
   when the caller asks for it (``allow_vision=True``, which today only
   ``read_document`` passes). Transcription is an explicit act, never ambient:
   the drafting record check consumes an existing cached transcription and
   never initiates one, so a scanned matter's draft path still opens only
   after a human deliberately read each scan.

``extract_text`` (the original) keeps its exact signature and behaviour and
NEVER reaches the network. It has callers outside this connector — the voice
corpus builders (``operator/bin/voice-fetch-corpus.py``,
``voice-survey-corpus.py``) and the pilot L2 driver — and a corpus builder that
silently started billing model calls per document would be a nasty surprise.
"""

from __future__ import annotations

import io
from dataclasses import dataclass

_PDF_MAGIC = b"%PDF"
_ZIP_MAGIC = b"PK\x03\x04"

# ---- The closed vocabularies -------------------------------------------------
# A caller branches on these, never on text truthiness: "empty string" is the
# ambiguity this module exists to remove.

METHOD_PYPDF = "pypdf"
METHOD_DOCX = "docx"
METHOD_PLAIN = "plain"
METHOD_VISION = "vision"
METHOD_VISION_CACHED = "vision_cached"
METHOD_NONE_SCANNED = "none_scanned"

METHODS = frozenset(
    {
        METHOD_PYPDF,
        METHOD_DOCX,
        METHOD_PLAIN,
        METHOD_VISION,
        METHOD_VISION_CACHED,
        METHOD_NONE_SCANNED,
    }
)

#: Why a scanned PDF produced no text. Closed set — never an exception message
#: (which can carry a URL or a credential fragment) and never model output.
REASON_NOT_ATTEMPTED = "not_attempted"  # caller passed allow_vision=False
REASON_NO_CREDENTIAL = "no_credential"
REASON_DISABLED = "disabled"
REASON_OVER_PAGE_CAP = "over_page_cap"
REASON_OVER_BYTE_CAP = "over_byte_cap"
REASON_API_ERROR = "api_error"
REASON_TRUNCATED = "truncated"
REASON_INCOMPLETE = "incomplete_transcription"

REASONS = frozenset(
    {
        REASON_NOT_ATTEMPTED,
        REASON_NO_CREDENTIAL,
        REASON_DISABLED,
        REASON_OVER_PAGE_CAP,
        REASON_OVER_BYTE_CAP,
        REASON_API_ERROR,
        REASON_TRUNCATED,
        REASON_INCOMPLETE,
    }
)

#: A PDF page carrying fewer than this many characters of extracted text is
#: paper, not a document: scanner artefacts, a stamped exhibit number, a fax
#: header. Averaged across the file so a scan bundle with one text-layer cover
#: sheet still reads as scanned.
SCANNED_CHARS_PER_PAGE = 15


class UnsupportedDocumentError(RuntimeError):
    """The file type has no text-extraction path. The message names the type so
    the agent can surface 'needs manual review' instead of guessing content."""


@dataclass(frozen=True)
class ExtractResult:
    """What came out, and by which road.

    ``text`` is empty for every non-success method. ``reason`` is set only for
    ``none_scanned`` and is one of :data:`REASONS`. ``pages`` is the PDF page
    count when known (it is what an attorney needs to judge an over-cap refusal)
    and None for every other type."""

    text: str
    method: str
    reason: str | None = None
    pages: int | None = None


def extract_text(blob: bytes, *, file_name: str = "", file_extension: str = "") -> str:
    """Extract plain text from a document blob. Detection is magic-bytes first
    (the extension metadata is client-supplied and can lie), extension second.

    MECHANICAL PATHS ONLY. This function never makes a network call and never
    consults the transcription cache — a scanned PDF returns ``""`` here exactly
    as it always has. Callers that want the vision fallback ask for it
    explicitly through :func:`extract_text_ex`."""
    return _extract(
        blob,
        file_name=file_name,
        file_extension=file_extension,
        allow_vision=False,
        allow_cache=False,
    ).text


def extract_text_ex(
    blob: bytes,
    *,
    file_name: str = "",
    file_extension: str = "",
    allow_vision: bool = False,
) -> ExtractResult:
    """Extract text and NAME the path it came from.

    Raises :class:`UnsupportedDocumentError` for a type with no extraction path,
    exactly as :func:`extract_text` does — an unsupported file is a refusal, not
    a result, and both callers already handle it that way.

    ``allow_vision`` gates the transcription of a scanned PDF. It is an explicit
    act: only ``read_document`` passes True. With it False a scanned PDF still
    consults the on-disk transcription cache (a read that costs nothing and
    cannot invent anything) and otherwise returns
    ``none_scanned(not_attempted)``."""
    return _extract(
        blob,
        file_name=file_name,
        file_extension=file_extension,
        allow_vision=allow_vision,
        allow_cache=True,
    )


def _extract(
    blob: bytes,
    *,
    file_name: str,
    file_extension: str,
    allow_vision: bool,
    allow_cache: bool,
) -> ExtractResult:
    """The one dispatch both public entry points share, so the mechanical
    detection can never drift between them."""
    ext = (file_extension or "").lower().lstrip(".")
    if blob.startswith(_PDF_MAGIC) or ext == "pdf":
        return _pdf_result(blob, allow_vision=allow_vision, allow_cache=allow_cache)
    if blob.startswith(_ZIP_MAGIC) and (ext in ("docx", "dotx", "docm", "dotm", "") or _looks_like_docx(blob)):
        return ExtractResult(_docx_text(blob), METHOD_DOCX)
    if ext in ("txt", "text", "md", "csv", "log", "eml"):
        return ExtractResult(blob.decode("utf-8", "replace"), METHOD_PLAIN)
    # Last resort: treat as text only when it plausibly IS text (no NUL bytes
    # in the first 4KB); otherwise refuse explicitly.
    if b"\x00" not in blob[:4096]:
        return ExtractResult(blob.decode("utf-8", "replace"), METHOD_PLAIN)
    raise UnsupportedDocumentError(
        f"no text-extraction path for {file_name or 'document'!r} "
        f"(extension {file_extension or 'unknown'!r}); needs manual review"
    )


def _pdf_result(blob: bytes, *, allow_vision: bool, allow_cache: bool) -> ExtractResult:
    """The PDF road: pypdf, then (only for paper) the cache, then vision."""
    text, pages = _pdf_text_and_pages(blob)
    if not _looks_scanned(text, pages):
        return ExtractResult(text, METHOD_PYPDF, pages=pages)
    if not allow_cache:
        return ExtractResult("", METHOD_NONE_SCANNED, reason=REASON_NOT_ATTEMPTED, pages=pages)

    from .extract_cache import cache_get, cache_put

    cached = cache_get(blob)
    if cached is not None:
        return ExtractResult(cached, METHOD_VISION_CACHED, pages=pages)
    if not allow_vision:
        return ExtractResult("", METHOD_NONE_SCANNED, reason=REASON_NOT_ATTEMPTED, pages=pages)

    from .vision import transcribe_pdf

    outcome = transcribe_pdf(blob, pages=pages)
    if outcome.reason is not None:
        return ExtractResult("", METHOD_NONE_SCANNED, reason=outcome.reason, pages=pages)
    cache_put(blob, outcome.text, pages=pages)
    return ExtractResult(outcome.text, METHOD_VISION, pages=pages)


def _looks_scanned(text: str, pages: int) -> bool:
    """True when the text layer yields ~nothing per page — i.e. this is paper."""
    if pages <= 0:
        return False
    return len(text.strip()) < SCANNED_CHARS_PER_PAGE * pages


def _looks_like_docx(blob: bytes) -> bool:
    # A DOCX is a zip whose early central-directory bytes name its parts.
    return b"[Content_Types].xml" in blob[:4096] or b"word/" in blob[:4096]


def _pdf_text_and_pages(blob: bytes) -> tuple[str, int]:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(blob))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - malformed PDFs must fail closed, not crash the server
        raise UnsupportedDocumentError(f"PDF could not be parsed: {exc}") from exc
    return "\n\n".join(p.strip() for p in pages if p.strip()), len(pages)


def _docx_text(blob: bytes) -> str:
    from docx import Document

    from .docx_base import _dotx_to_docx_bytes, _is_dotx

    # A Word TEMPLATE (.dotx/.dotm) is a .docx with one content-type string
    # changed; python-docx refuses it as-is. A firm's letterhead template filed
    # on a matter must extract like any other document, not poison every
    # record check on that matter.
    if _is_dotx(blob):
        blob = _dotx_to_docx_bytes(blob)
    try:
        doc = Document(io.BytesIO(blob))
    except Exception as exc:  # noqa: BLE001 - malformed DOCX must fail closed, not crash the server
        raise UnsupportedDocumentError(f"DOCX could not be parsed: {exc}") from exc
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)
