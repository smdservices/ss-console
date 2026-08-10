"""Markdown skeleton -> .docx rendering for the ``render_docx_template`` tool.

Two halves, both deliberately mechanical:

1. **The content gate** (:func:`check_template_content`) — a pure function over
   the markdown, unit-testable without a network or a document object. It
   REFUSES; it never strips, repairs, or downgrades. What it refuses is not a
   style opinion, it is the four ways a *template* stops being a template:

   - **Case content outside a ``{{...}}`` marker.** A template carries
     structure, not a case, and content that reaches a template reaches every
     future matter the template is filled for. The gate refuses case-content
     SHAPES, not digits: a template is *full* of legitimate numbers, because
     statutory citations, code sections, and statutory periods ARE template
     structure. "Code of Civil Procedure section 999", "not fewer than 30 days",
     "CCP 2030.060(f)", "Vehicle Code section 22350", and a reported case year
     all belong in a skeleton and all pass. Four shapes do not:

     1. **a date** — ``2024-03-01``, ``3/1/2024``, ``March 1, 2024``, ``Mar. 1 2024``
     2. **a dollar figure** — ``$`` followed by digits, commas and decimals optional
     3. **an identifier** — a letter-prefixed run (``ZZ-9999-0001``), a case
        number (``2026-PI-102``), or a hyphenated numeric range (bates)
     4. **a claim or bates number** — a bare run of five or more digits

     Shape 4 is the one that can collide with legitimate structure: California
     code sections run to five digits (Vehicle Code 22350, 21703, 21453). A bare
     long run is therefore case content only when nothing cites it as law, so a
     statutory-citation span (``section(s) N``, ``§ N``, including the
     comma-joined lists these come in) exempts shape 4 and no other. A hyphen is
     not a citation joiner: ``sections 000123-000456`` is a bates range wearing a
     citation's clothes and stays refused.

     Digits INSIDE a marker are always fine: a marker names its own source, so
     ``{{FILL: date of loss | traffic collision report}}`` and
     ``{{NOT IN RECORD: subpart check under CCP 2030.060(f)}}`` are structure,
     not content. Text inside an HTML comment is exempt from this rule for the
     same reason it is refused by the next one: the comment is already a
     violation and is destined for deletion, so scanning inside it would report
     the same removal twice and overstate how many things need fixing. It can
     never turn a refusal into a pass.
   - **Malformed marker syntax** — an unbalanced ``{{`` or ``}}``, or an empty
     marker. A marker that does not close is not a marker; it is a sentence
     fragment that a filler will read as prose and quietly answer.
   - **An em dash.** Banned in shipped copy by house style, and by drafting
     discipline rule 7 ("No em dashes") for every draft this template produces.
   - **An HTML comment.** Drafting gate 9 (visible-delta rule): a reservation or
     a divergence note must survive rendering as body text. ``<!-- ... -->``
     vanishes into a .docx as nothing at all, so an attorney reviewing the
     rendered document never sees the thing that was reserved. This is the whole
     failure mode the gate exists to prevent, and it is invisible by
     construction, which is why it is mechanical rather than advisory.

   Every violation in the document is reported, not just the first — a caller
   that fixes one and resubmits four times has been told the truth four times
   and still does not know what is wrong.

2. **The renderer** (:func:`render_markdown_to_docx`) — a small, honest subset of
   markdown. Headings (``#``/``##``/``###``), paragraphs, ``-``/``*`` bullets,
   and ``**bold**``/``*italic*`` inline. Everything else renders as plain
   paragraph text: a construct this renderer does not understand is shown to the
   reader verbatim, NEVER dropped. Markers are emitted as literal, unstyled runs
   and are never touched by the inline-emphasis pass, so a marker containing an
   asterisk or an underscore survives byte-for-byte and stays render-visible.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

_EM_DASH = "—"
_HTML_COMMENT_OPEN = "<!--"
_HTML_COMMENT_CLOSE = "-->"

_MONTHS = "Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec"

# The case-content shapes. Each carries the plain name that goes in the refusal,
# so the caller is told WHAT it looks like it left in the template rather than
# "there is a digit here". Statutory citations, code sections, and statutory
# periods match none of these by design — they are template structure.
_CASE_CONTENT_SHAPES: tuple[tuple[str, re.Pattern[str], bool], ...] = (
    ("a date", re.compile(r"\b\d{4}-\d{2}-\d{2}\b"), False),
    ("a date", re.compile(r"\b\d{1,2}/\d{1,2}/\d{2,4}\b"), False),
    (
        "a date",
        re.compile(
            rf"\b(?:{_MONTHS})[a-z]*\.?\s+\d{{1,2}}(?:st|nd|rd|th)?,?\s+\d{{4}}\b"
        ),
        False,
    ),
    ("a dollar figure", re.compile(r"\$\s?\d[\d,]*(?:\.\d+)?"), False),
    # ZZ-9999-0001 (the sentinel class) and any letter-prefixed identifier.
    ("an identifier", re.compile(r"\b[A-Z]{2,}-\d{3,}(?:-\d+)*\b"), False),
    # 2026-PI-102 — the matter-number shape; digit runs too short for the
    # bare-run rule, which is exactly why it needs its own.
    ("a case number", re.compile(r"\b\d{4}-[A-Z]{1,4}-\d+\b"), False),
    ("a bates or identifier range", re.compile(r"\b\d{3,}[-–]\d{3,}\b"), False),
    # The one shape a statutory citation can collide with: California code
    # sections run to five digits (Vehicle Code 22350, 21703, 21453), so a bare
    # long run is a claim number ONLY when nothing cites it as law. Hence the
    # citation exemption below, which applies to this shape and no other.
    ("a claim or bates number", re.compile(r"\b\d{5,}\b"), True),
)

# A statutory citation, including the comma-and-"and"-joined lists these come in
# ("sections 22350, 21703, 21801, 22107, and 21453"; "sections 999 through
# 999.5"). Hyphens are deliberately NOT joiners: "sections 000123-000456" is a
# bates range wearing a citation's clothes, and it stays refused.
_CITATION_JOINER = r"(?:\s*(?:,|;|and|or|through|to)\s*)+"
_STATUTORY_CITATION_RE = re.compile(
    rf"(?:§§?|\bsections?\b)\s*\d[\d.]*(?:{_CITATION_JOINER}\d[\d.]*)*",
    re.IGNORECASE,
)

# Well-formed markers, for the render pass. The gate runs first, so by the time
# the renderer sees the text every ``{{`` has a ``}}``; non-greedy so adjacent
# markers on one line stay separate.
_MARKER_RE = re.compile(r"\{\{.*?\}\}", re.DOTALL)

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
# Split on emphasis spans, keeping them (capturing group). ``[^*]+`` keeps the
# match tight so ``**a** and **b**`` yields two bold runs, not one.
_EMPHASIS_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")

_SNIPPET_CHARS = 100


@dataclass(frozen=True)
class Violation:
    """One refusal reason, located. ``rule`` is a stable machine-readable slug so
    a caller (or a test) can assert on the class of failure rather than on
    prose."""

    rule: str
    line: int
    detail: str

    def __str__(self) -> str:
        return f"line {self.line}: [{self.rule}] {self.detail}"


class TemplateContentRefused(RuntimeError):
    """The skeleton is not renderable as a template. Carries EVERY violation in
    ``.violations`` — the tool surfaces the whole list as a refusal, and nothing
    is uploaded."""

    def __init__(self, violations: list[Violation]) -> None:
        self.violations = list(violations)
        body = "\n".join(f"  - {v}" for v in self.violations)
        super().__init__(
            f"refusing to render: {len(self.violations)} template-content "
            f"violation(s). Fix the source markdown; nothing was uploaded.\n{body}"
        )


# ---- The gate --------------------------------------------------------------


def check_template_content(markdown: str) -> None:
    """Raise :class:`TemplateContentRefused` if the skeleton is not a template.
    Returns None (and nothing else) when it is clean."""
    violations = find_violations(markdown)
    if violations:
        raise TemplateContentRefused(violations)


def find_violations(markdown: str) -> list[Violation]:
    """Every content violation in ``markdown``, ordered by position. Pure — no
    I/O, no document object, no network — so the refusal contract is testable on
    its own."""
    line_starts = _line_starts(markdown)
    spans, violations = _scan_markers(markdown, line_starts)
    exempt = spans + _comment_spans(markdown)
    violations.extend(_case_content_violations(markdown, exempt, line_starts))
    violations.extend(_literal_violations(markdown, line_starts))
    return sorted(violations, key=lambda v: (v.line, v.rule))


def _line_starts(text: str) -> list[int]:
    starts = [0]
    for i, ch in enumerate(text):
        if ch == "\n":
            starts.append(i + 1)
    return starts


def _line_of(line_starts: list[int], index: int) -> int:
    """1-based line number for a character offset (binary search)."""
    import bisect

    return bisect.bisect_right(line_starts, index)


def _snippet(text: str, line_starts: list[int], line: int) -> str:
    start = line_starts[line - 1]
    end = text.find("\n", start)
    if end == -1:
        end = len(text)
    body = text[start:end].strip()
    if len(body) > _SNIPPET_CHARS:
        body = body[:_SNIPPET_CHARS] + "..."
    return body


def _scan_markers(
    text: str, line_starts: list[int]
) -> tuple[list[tuple[int, int]], list[Violation]]:
    """Find every well-formed ``{{...}}`` span and every way the marker syntax is
    broken. Spans are (start, end) half-open and INCLUDE the delimiters, so a
    digit anywhere between ``{{`` and ``}}`` counts as inside a marker.

    An unterminated ``{{`` yields no span, deliberately: pretending the rest of
    the document is one giant marker would hide every digit after it behind the
    very defect being reported."""
    spans: list[tuple[int, int]] = []
    violations: list[Violation] = []
    i = 0
    n = len(text)
    while i < n:
        open_at = text.find("{{", i)
        close_at = text.find("}}", i)
        if open_at == -1 and close_at == -1:
            break
        if close_at != -1 and (open_at == -1 or close_at < open_at):
            violations.append(
                Violation(
                    "marker-syntax",
                    _line_of(line_starts, close_at),
                    "closing '}}' with no matching '{{'",
                )
            )
            i = close_at + 2
            continue
        end = text.find("}}", open_at + 2)
        if end == -1:
            violations.append(
                Violation(
                    "marker-syntax",
                    _line_of(line_starts, open_at),
                    "opening '{{' is never closed",
                )
            )
            i = open_at + 2
            continue
        inner = text[open_at + 2 : end]
        if "{{" in inner:
            violations.append(
                Violation(
                    "marker-syntax",
                    _line_of(line_starts, open_at),
                    "nested '{{' inside a marker",
                )
            )
        if not inner.strip():
            violations.append(
                Violation(
                    "marker-syntax",
                    _line_of(line_starts, open_at),
                    "empty marker '{{}}' names no source",
                )
            )
        spans.append((open_at, end + 2))
        i = end + 2
    return spans, violations


def _comment_spans(text: str) -> list[tuple[int, int]]:
    """HTML-comment spans, exempt from the case-content scan. An unterminated
    ``<!--`` runs to the end of the text; that can never turn a refusal into a
    pass, because the comment itself is already a violation."""
    spans: list[tuple[int, int]] = []
    start = text.find(_HTML_COMMENT_OPEN)
    while start != -1:
        close = text.find(_HTML_COMMENT_CLOSE, start + len(_HTML_COMMENT_OPEN))
        end = len(text) if close == -1 else close + len(_HTML_COMMENT_CLOSE)
        spans.append((start, end))
        start = text.find(_HTML_COMMENT_OPEN, end)
    return spans


def _in_spans(index: int, spans: list[tuple[int, int]]) -> bool:
    return any(start <= index < end for start, end in spans)


def _case_content_violations(
    text: str, exempt: list[tuple[int, int]], line_starts: list[int]
) -> list[Violation]:
    """Case-content SHAPES outside a marker: dates, dollar figures, identifiers,
    and long bare digit runs. Statutory citations and periods are structure and
    pass — see the module docstring.

    One violation per offending LINE, not per match: a table row of figures is
    one problem to fix, and 40 entries would bury the other rules."""
    cited = [(m.start(), m.end()) for m in _STATUTORY_CITATION_RE.finditer(text)]
    out: list[Violation] = []
    seen: set[int] = set()
    for shape, pattern, honors_citation in _CASE_CONTENT_SHAPES:
        for match in pattern.finditer(text):
            if _in_spans(match.start(), exempt):
                continue
            if honors_citation and _in_spans(match.start(), cited):
                continue
            line = _line_of(line_starts, match.start())
            if line in seen:
                continue
            seen.add(line)
            out.append(
                Violation(
                    "case-content",
                    line,
                    f"{shape} ({match.group(0)!r}) in a template reaches every "
                    "matter the template is filled for. Put it in a "
                    "{{FILL: ... | source}} marker: "
                    f"{_snippet(text, line_starts, line)!r}",
                )
            )
    return sorted(out, key=lambda v: v.line)


def _literal_violations(text: str, line_starts: list[int]) -> list[Violation]:
    """Em dashes (house style + drafting rule 7) and HTML comments (drafting gate
    9: a reservation that vanishes on render was never reserved)."""
    out: list[Violation] = []
    seen_em: set[int] = set()
    start = text.find(_EM_DASH)
    while start != -1:
        line = _line_of(line_starts, start)
        if line not in seen_em:
            seen_em.add(line)
            out.append(
                Violation(
                    "em-dash",
                    line,
                    "em dash is banned in shipped copy and in every draft this "
                    f"template produces: {_snippet(text, line_starts, line)!r}",
                )
            )
        start = text.find(_EM_DASH, start + 1)
    start = text.find(_HTML_COMMENT_OPEN)
    while start != -1:
        out.append(
            Violation(
                "html-comment",
                _line_of(line_starts, start),
                "an HTML comment renders as nothing: guidance, reservations, and "
                "divergence notes must be render-visible body text",
            )
        )
        start = text.find(_HTML_COMMENT_OPEN, start + 1)
    return out


# ---- The renderer ----------------------------------------------------------


def render_markdown_to_docx(markdown: str) -> bytes:
    """Render gated skeleton markdown to .docx bytes.

    Supported: ``#``/``##``/``###`` -> Heading 1/2/3; ``-``/``*`` bullets ->
    List Bullet; ``**bold**`` / ``*italic*`` inline; blank lines separate
    paragraphs. Anything else (deeper headings, tables, rules, links, code
    fences) renders as plain paragraph text with its markdown characters intact:
    the reader sees it, which is the only failure mode acceptable in a document
    an attorney reviews.

    Markers are emitted as their own literal runs with no formatting applied, so
    a marker is never restyled, hidden, or split by the emphasis pass."""
    from docx import Document

    doc = Document()
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        heading = _HEADING_RE.match(line)
        if heading:
            para = doc.add_paragraph(style=f"Heading {len(heading.group(1))}")
            _add_runs(para, heading.group(2).strip())
            continue
        bullet = _BULLET_RE.match(line)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            _add_runs(para, bullet.group(1).strip())
            continue
        _add_runs(doc.add_paragraph(), line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs(paragraph, text: str) -> None:
    """Emit ``text`` as runs, with marker spans held out of emphasis parsing so
    they survive verbatim."""
    pos = 0
    for match in _MARKER_RE.finditer(text):
        _add_formatted(paragraph, text[pos : match.start()])
        paragraph.add_run(match.group(0))  # literal, unstyled, render-visible
        pos = match.end()
    _add_formatted(paragraph, text[pos:])


def _add_formatted(paragraph, text: str) -> None:
    if not text:
        return
    for part in _EMPHASIS_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)
