"""Base-document handling for the format renderer: open the firm's template (or
the stock starter), clear its body, keep everything else; the named-style
contract and the product-default look applied when a template lacks a style.

Split out of ``docx_format`` for size only; ``docx_format`` re-exports what the
tools and tests use. See that module's docstring for the doctrine.
"""

from __future__ import annotations

import io
import zipfile

from .docx_format_types import DEFAULT_FONT, DEFAULT_SIZE_PT, NAMED_STYLES, FormatRefused, FormatReport

# ---- Base document -------------------------------------------------------------

_DOTX_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
_DOCX_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"


def _dotx_to_docx_bytes(blob: bytes) -> bytes:
    """python-docx rejects a ``.dotx`` (template content type). The only
    difference that matters is one content-type string; rewrite it so the
    firm's template opens as a document."""
    src = zipfile.ZipFile(io.BytesIO(blob))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(_DOTX_MAIN.encode(), _DOCX_MAIN.encode())
            dst.writestr(item, data)
    return out.getvalue()


def _is_dotx(blob: bytes) -> bool:
    try:
        with zipfile.ZipFile(io.BytesIO(blob)) as z:
            return _DOTX_MAIN.encode() in z.read("[Content_Types].xml")
    except (zipfile.BadZipFile, KeyError):
        return False


def open_as_base(blob: bytes | None, report: FormatReport):
    """Open the firm's template (or the stock base) and clear its body, keeping
    the final ``w:sectPr`` so page setup, headers and footers survive.

    A multi-section document is refused: clearing the body would keep only the
    LAST section's page setup and silently drop the rest, and that is exactly
    the wrong letterhead on a client letter."""
    from docx import Document
    from docx.oxml.ns import qn

    if blob is None:
        # The STARTER: python-docx's stock base with the authored-standard
        # docDefaults and the named-style contract defined, self-described in
        # its properties so a firm opening it in Word knows what it is and that
        # the styles are theirs to edit. Filed into a library via
        # render_docx_template(document_class=...), it becomes provenance (b).
        doc = Document()
        _rewrite_doc_defaults(doc)
        _ensure_named_styles(doc, report)
        cp = doc.core_properties
        cp.author = "SMD Operator"
        cp.comments = (
            "SMD starter template. Typography lives in this file's styles "
            "(SMD Body, SMD Item Label, SMD Item Text, SMD Heading 1-3, SMD Caption, "
            "SMD Signature): edit them in Word and the next draft follows."
        )
        return doc
    if _is_dotx(blob):
        blob = _dotx_to_docx_bytes(blob)
        report.notes.append("base was a .dotx; opened as a document")
    doc = Document(io.BytesIO(blob))
    body = doc.element.body
    sects = body.findall(qn("w:sectPr"))
    if len(doc.sections) > 1 or len(sects) > 1:
        raise FormatRefused(
            f"the base template has {len(doc.sections)} sections; the renderer keeps one "
            "section's page setup and would silently drop the rest. Use a single-section "
            "template (a different first page is fine: that is a header setting, not a section)."
        )
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    report.base_header_footer_text = _header_footer_text(doc)
    return doc


def _header_footer_text(doc) -> list[str]:
    out: list[str] = []
    for section in doc.sections:
        parts = [section.header, section.footer]
        # Only read first-page parts when the section declares them; touching
        # them otherwise would CREATE empty parts in the firm's file.
        if section.different_first_page_header_footer:
            parts += [section.first_page_header, section.first_page_footer]
        for part in parts:
            if part.is_linked_to_previous:
                continue
            for p in part.paragraphs:
                t = p.text.strip()
                if t:
                    out.append(t)
    return out


def _rewrite_doc_defaults(doc) -> None:
    """python-docx's stock base carries Calibri 11 with 1.15 line spacing and
    space-after in docDefaults; every paragraph inherits it. The starter look is
    the authored standard instead."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    for tag in ("w:rPrDefault", "w:pPrDefault"):
        old = dd.find(qn(tag))
        if old is not None:
            dd.remove(old)
    rpr_default = OxmlElement("w:rPrDefault")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), DEFAULT_FONT)
    rpr.append(fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(DEFAULT_SIZE_PT * 2))
    rpr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), str(DEFAULT_SIZE_PT * 2))
    rpr.append(szcs)
    rpr_default.append(rpr)
    ppr_default = OxmlElement("w:pPrDefault")
    ppr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    ppr_default.append(ppr)
    dd.append(rpr_default)
    dd.append(ppr_default)
    # The stock Normal style also pins Calibri via the theme; make Normal follow.
    normal = doc.styles["Normal"]
    _set_font(normal.font, DEFAULT_FONT, DEFAULT_SIZE_PT)
    normal.paragraph_format.space_after = None
    normal.paragraph_format.line_spacing = None


def _set_font(font, name: str, size_pt: float | None) -> None:
    """``font.name`` sets only ascii/hAnsi; Word reads eastAsia/cs too, so set
    all four or a theme font leaks through on some runs."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    font.name = name
    rpr = font.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is not None:
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            fonts.set(qn(attr), name)
    if size_pt is not None:
        font.size = Pt(size_pt)


def has_style(doc, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def usable_paragraph_style(doc, name: str) -> str | None:
    """The name if the base defines it as a paragraph style we can safely apply,
    else None with nothing raised.

    ``has_style`` alone is not enough once we delegate to a FIRM's own style
    names, because three real shapes make a present style unusable:

    1. **Not a paragraph style.** python-docx raises ``ValueError`` (not
       ``KeyError``) when a character or table style is assigned to a paragraph,
       so a firm with a character style named "Heading 2" would kill the whole
       render rather than degrade. The hazard predates delegation: ``_bullet``
       and ``_table`` test ``has_style`` and assign directly.
    2. **Outline numbering.** Firms routinely link Heading 1/2 to a multilevel
       list. Our drafting grammar has the model write the numeral itself
       ("## I. Introduction"), so a numbered base style yields "1.1 I.
       Introduction" on a filed document. Numbering is inherited, so the
       ``base_style`` chain is walked, not just the style itself.
    3. **Latent styles.** A .docx saved by Word carries ``w:style`` elements
       only for styles in use; the rest live in ``w:latentStyles`` and are
       invisible to ``doc.styles``. Nothing to do about it here beyond not
       claiming the style exists — which is why the caller reports the gap as
       something the firm can fix in Word.
    """
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    try:
        style = doc.styles[name]
    except KeyError:
        return None
    if getattr(style, "type", None) != WD_STYLE_TYPE.PARAGRAPH:
        return None
    seen: set[int] = set()
    node = style
    while node is not None and id(node) not in seen:
        seen.add(id(node))
        element = getattr(node, "element", None)
        if element is not None and element.find(qn("w:pPr")) is not None:
            if element.find(qn("w:pPr")).find(qn("w:numPr")) is not None:
                return None
        node = getattr(node, "base_style", None)
    return name


def _ensure_named_styles(doc, report: FormatReport) -> None:
    """Define the named-style contract in a document WE own (the stock base /
    a starter). Never called on a firm's file: their styles are theirs."""
    from docx.enum.style import WD_STYLE_TYPE

    for name in NAMED_STYLES:
        if has_style(doc, name):
            continue
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        _apply_default_style_spec(style, name)
    report.notes.append("named styles defined on the starter base")


def _apply_default_style_spec(style, name: str) -> None:
    """The product default look for each named style (the starter)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    f, pf = style.font, style.paragraph_format
    _set_font(f, DEFAULT_FONT, DEFAULT_SIZE_PT)
    if name == "SMD Item Label":
        f.bold = True
        f.underline = True
        f.all_caps = True
        pf.space_before = Pt(12)
    elif name == "SMD Item Text":
        pf.first_line_indent = Inches(0.5)
        pf.line_spacing = 2.0
        pf.space_after = Pt(12)
    elif name == "SMD Heading 1":
        f.bold = True
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        pf.keep_with_next = True
    elif name == "SMD Heading 2":
        f.bold = True
        f.underline = True
        pf.left_indent = Inches(0.5)
        pf.space_before = Pt(6)
        pf.keep_with_next = True
    elif name == "SMD Heading 3":
        f.bold = True
        pf.left_indent = Inches(1.0)
        pf.keep_with_next = True
    elif name == "SMD Caption":
        pf.space_after = Pt(0)
    elif name == "SMD Signature":
        pf.left_indent = Inches(3.5)
        pf.space_before = Pt(24)
    # SMD Body: Normal + the default font; nothing else.
