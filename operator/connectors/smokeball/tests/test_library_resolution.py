"""library.resolve_template + the render tools' ``document_class`` path.

The resolution is the tool's, not the model's: the same customer.yaml, the
same matter, the same folder, the same file name, every time. "Not resolved"
is a reported outcome (the draft still files, on the starter base), never a
refusal and never silent.
"""

from __future__ import annotations

import hashlib
import io
import textwrap
import zipfile

import httpx
import pytest
from docx import Document

from smokeball_connector import server
from smokeball_connector.library import (
    CUSTOMER_YAML_ENV,
    DEFAULT_FOLDER_NAME,
    LOOKUP_FAILED,
    LOOKUP_FOUND,
    LOOKUP_NOT_FOUND,
    OPERATOR_LIBRARY_NUMBER,
    LibraryConfig,
    NotResolved,
    ResolvedTemplate,
    is_library_file,
    load_library_config,
    lookup_matter,
    resolve_template,
)
from smokeball_connector.render import render_markdown_to_docx

from .test_render_document import DISCOVERY_MD, make_firm_template

_UPLOAD_URL = "https://s3.example.com/apiuploads/draft-2?X-Amz-Signature=deadbeef"
_DOWNLOAD_URL = "https://s3.example.com/apidownloads/tpl-1?X-Amz-Signature=cafebabe"


# ---- config -------------------------------------------------------------------------


def _write_yaml(tmp_path, body: str) -> str:
    p = tmp_path / "customer.yaml"
    p.write_text(textwrap.dedent(body))
    return str(p)


def test_missing_customer_yaml_is_not_authored_with_a_reason(tmp_path) -> None:
    cfg = load_library_config(str(tmp_path / "nope.yaml"))
    assert cfg.authored is False
    assert "not readable" in cfg.source


def test_block_without_a_number_falls_back_to_the_operator_library_convention(tmp_path) -> None:
    """CHANGED by ss-console#2536, deliberately. Before the Operator could create
    its own matter, a block with no ``matter_number`` resolved to nothing and the
    report said "not authored". Now the number the Operator would create under is
    the last fallback, so a seat resolves its library before anybody has authored
    a number, and ``fallback_number`` carries the distinction into the report:
    "the firm's library matter is missing" and "the matter we would have created
    is not there yet" are different sentences to an admin."""
    path = _write_yaml(tmp_path, """
        self_initiation:
          document_library:
            matter_hint: '2026-OPS-001 (the internal operations matter)'
            folder_name: 'Document Library'
    """)
    cfg = load_library_config(path)
    assert cfg.authored is True
    assert cfg.matter_number == OPERATOR_LIBRARY_NUMBER
    assert cfg.fallback_number is True
    assert OPERATOR_LIBRARY_NUMBER in cfg.source
    assert cfg.folder_name == "Document Library"


def test_the_authored_operator_matter_number_wins_over_the_convention(tmp_path) -> None:
    path = _write_yaml(tmp_path, """
        self_initiation:
          document_library:
            folder_name: 'Document Library'
            operator_matter:
              number: '2026-OPS-LIBRARY'
              description: 'Operator Library'
              client_contact_id: 'c-1'
              matter_type_id: 't-1'
    """)
    cfg = load_library_config(path)
    assert cfg.matter_number == "2026-OPS-LIBRARY"
    assert cfg.fallback_number is False


def test_an_authored_matter_number_wins_over_the_operator_matter_block(tmp_path) -> None:
    """The firm's own library matter is the first answer, always. A seat that
    already keeps templates somewhere does not get moved by this feature."""
    path = _write_yaml(tmp_path, """
        self_initiation:
          document_library:
            matter_number: '2026-OPS-001'
            folder_name: 'Document Library'
            operator_matter:
              number: 'OPS-OPERATOR-LIBRARY'
              description: 'Operator Library'
              client_contact_id: 'c-1'
              matter_type_id: 't-1'
    """)
    cfg = load_library_config(path)
    assert cfg.matter_number == "2026-OPS-001"
    assert cfg.fallback_number is False


def test_an_absent_block_is_still_not_authored(tmp_path) -> None:
    """The fallback fills a MISSING FIELD in a block the firm wrote. It does not
    invent a library for a seat that never asked for one."""
    path = _write_yaml(tmp_path, "self_initiation:\n  sequence: []\n")
    cfg = load_library_config(path)
    assert cfg.authored is False
    assert cfg.matter_number is None
    assert "not authored" in cfg.source


def test_a_block_without_a_folder_name_uses_the_proposed_default(tmp_path) -> None:
    path = _write_yaml(tmp_path, """
        self_initiation:
          document_library:
            matter_number: '2026-OPS-001'
    """)
    cfg = load_library_config(path)
    assert cfg.folder_name == DEFAULT_FOLDER_NAME


def test_authored_block_and_template_name_convention_and_override(tmp_path) -> None:
    path = _write_yaml(tmp_path, """
        self_initiation:
          document_library:
            matter_number: '2026-OPS-001'
            folder_name: 'Document Library'
            templates:
              letter: 'Firm Letterhead'
    """)
    cfg = load_library_config(path)
    assert cfg.authored is True and cfg.matter_number == "2026-OPS-001"
    assert cfg.template_name("discovery_set") == "Template - Discovery Set.docx"
    assert cfg.template_name("letter") == "Firm Letterhead.docx"


def test_env_override_points_the_loader_at_a_file(tmp_path, monkeypatch) -> None:
    path = _write_yaml(tmp_path, "self_initiation:\n  document_library:\n    matter_number: 'M-1'\n    folder_name: 'Lib'\n")
    monkeypatch.setenv(CUSTOMER_YAML_ENV, path)
    assert load_library_config().authored is True


# ---- resolution against a fake client ------------------------------------------------


class _FakeClient:
    def __init__(self, *, matters, folders, files, blob: bytes = b"PK-template") -> None:
        self.matters, self.folders, self.files, self.blob = matters, folders, files, blob
        self.downloads: list[tuple[str, str]] = []

    def get(self, path: str, **params):
        if path == "/matters":
            return {"value": self.matters}
        if path.endswith("/documents/folders"):
            return {"value": self.folders}
        if path.endswith("/documents/files"):
            return {"value": self.files}
        raise AssertionError(path)

    def download_file(self, matter_id: str, file_id: str):
        self.downloads.append((matter_id, file_id))
        return {"name": "x"}, self.blob


_CFG = LibraryConfig(authored=True, matter_number="2026-OPS-001", folder_name="Document Library", source="test")


def test_resolves_the_class_template_in_the_library_folder() -> None:
    client = _FakeClient(
        matters=[{"id": "m-ops", "number": "2026-OPS-001"}, {"id": "m-other", "number": "2026-OPS-0010"}],
        folders=[{"id": "f-lib", "name": "Document Library"}],
        files=[
            {"id": "old", "name": "Template - Discovery Set.docx", "folderId": "f-lib", "dateCreated": "2026-01-01"},
            {"id": "new", "name": "template - discovery set.docx", "folderId": "f-lib", "dateCreated": "2026-08-01"},
            {"id": "root", "name": "Template - Discovery Set.docx", "folderId": None, "dateCreated": "2026-09-01"},
        ],
    )
    out = resolve_template(client, _CFG, "discovery_set")
    assert isinstance(out, ResolvedTemplate)
    assert out.file_id == "new"  # newest IN THE FOLDER wins; the root copy is ignored when folder matches exist
    assert out.matter_id == "m-ops" and out.folder_id == "f-lib"
    assert client.downloads == [("m-ops", "new")]


def test_not_authored_is_reported_not_raised() -> None:
    out = resolve_template(_FakeClient(matters=[], folders=[], files=[]), LibraryConfig(authored=False, source="x"), "memo")
    assert isinstance(out, NotResolved) and "not authored" in out.reason


def test_missing_matter_folder_and_file_each_name_their_reason() -> None:
    no_matter = _FakeClient(matters=[{"id": "m", "number": "OTHER"}], folders=[], files=[])
    assert "not found" in resolve_template(no_matter, _CFG, "memo").reason
    no_file = _FakeClient(matters=[{"id": "m", "number": "2026-OPS-001"}], folders=[{"id": "f", "name": "Document Library"}], files=[])
    out = resolve_template(no_file, _CFG, "memo")
    assert isinstance(out, NotResolved) and "Template - Memo.docx" in out.reason and out.folder_id == "f"
    no_folder = _FakeClient(matters=[{"id": "m", "number": "2026-OPS-001"}], folders=[], files=[])
    assert "folder 'Document Library' not found" in resolve_template(no_folder, _CFG, "memo").reason


_LIVE_ENTRY_IN_FOLDER = {
    # Pinned to the wire shape observed on the pilot tenant 2026-08-19
    # (vfy_01M0DTM2EGQZP9FZTM53S17CJ7). Not a guess.
    "href": "https://stagingapi.smokeball.com/matters/3c19.../documents/files/1410...",
    "id": "14105616-27d2-45f3-b212-40a12714060a",
    "versionId": "14105616-27d2-45f3-b212-40a12714060a36F5...",
    "folder": {"id": "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8", "href": "https://stagingapi.smokeball.com/matters/3c19.../documents/folders/9898..."},
    "matter": {"id": "3c191bed-cdda-48b9-a6ed-a51a349f3f94", "href": "https://stagingapi.smokeball.com/matters/3c19..."},
    "name": "Template - Demand Letter (Policy Limits)",
    "fileExtension": ".docx",
    "dateCreated": "2026-08-11T20:11:06.264423Z",
    "sizeBytes": 39150,
    "additionalData": {},
    "isUploaded": True,
    "isDeleted": False,
}
_LIVE_ENTRY_AT_ROOT = {
    "id": "df59b111-1bb2-4af0-b7e0-69454ff080be",
    "matter": {"id": "3c191bed-cdda-48b9-a6ed-a51a349f3f94"},
    "name": "Operator Self-Test Report 2026-08-11",
    "fileExtension": ".docx",
    "dateCreated": "2026-08-11T18:53:04.899418Z",
}


# The folder listing is a TREE, pinned to the payload observed on the pilot
# tenant 2026-08-20 (probe F). ``value`` holds one root node with `folders` and
# `files`; the root has no name. A flat read finds nothing, and that is exactly
# how the firm's template silently failed to resolve on the first live run.
_LIVE_FOLDER_LISTING = {
    "href": "https://stagingapi.smokeball.com/matters/3c19.../documents/folders",
    "offset": 0,
    "limit": 50,
    "size": 2,
    "value": [
        {
            "folders": [
                {
                    "href": "https://stagingapi.smokeball.com/matters/3c19.../documents/folders/9898...",
                    "id": "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8",
                    "name": "Document Library",
                }
            ],
            "files": [
                {
                    "id": "df59b111-1bb2-4af0-b7e0-69454ff080be",
                    "name": "Operator Self-Test Report 2026-08-11",
                }
            ],
        }
    ],
}


class _TreeClient(_FakeClient):
    """A client whose folder listing is the observed tree."""

    def get(self, path: str, **params):
        if path.endswith("/documents/folders"):
            return _LIVE_FOLDER_LISTING
        return super().get(path, **params)


def test_stored_name_carries_no_extension_so_the_match_is_extension_agnostic() -> None:
    """OBSERVED (probe H, 2026-08-20): Smokeball stores `name` WITHOUT the
    extension and `fileExtension` separately. An exact full-name match never
    matches anything, which is how the firm's template stayed unresolvable
    even after the folder was found."""
    from smokeball_connector.library import name_matches

    stored = {"name": "Template - Discovery Set", "fileExtension": ".docx"}
    assert name_matches(stored, "Template - Discovery Set.docx")
    assert name_matches(stored, "Template - Discovery Set")
    assert not name_matches(stored, "Template - Demand Letter.docx")
    # an authored override typed either way resolves
    authored = {"name": "Template - Demand Letter (Policy Limits)", "fileExtension": ".docx"}
    assert name_matches(authored, "Template - Demand Letter (Policy Limits).docx")
    assert name_matches(authored, "Template - Demand Letter (Policy Limits)")
    # a file whose name really does carry the extension still matches
    assert name_matches({"name": "Template - Memo.docx"}, "Template - Memo.docx")


def test_resolution_against_the_stored_name_shape() -> None:
    client = _TreeClient(
        matters=[{"id": "m-ops", "number": "2026-OPS-001"}],
        folders=[],
        files=[
            {
                "id": "tpl",
                "name": "Template - Discovery Set",
                "fileExtension": ".docx",
                "folder": {"id": "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8"},
                "dateCreated": "2026-08-20",
            }
        ],
        blob=b"PK-firm-template",
    )
    out = resolve_template(client, _CFG, "discovery_set")
    assert isinstance(out, ResolvedTemplate) and out.file_id == "tpl"


def test_folder_listing_tree_is_walked_not_read_flat() -> None:
    from smokeball_connector.library import find_folder_id

    client = _TreeClient(matters=[], folders=[], files=[])
    assert find_folder_id(client, "m-ops", "Document Library") == "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8"
    assert find_folder_id(client, "m-ops", "document library") == "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8"
    assert find_folder_id(client, "m-ops", "Nope") is None


def test_nested_subfolders_are_reachable() -> None:
    from smokeball_connector.library import _walk_folders

    tree = [{"folders": [{"id": "a", "name": "A", "folders": [{"id": "b", "name": "B"}]}]}]
    assert {f["id"] for f in _walk_folders(tree)} == {"a", "b"}


def test_resolution_end_to_end_against_the_observed_tree_and_entry_shapes() -> None:
    """The whole chain on the shapes the tenant actually returns: matter by
    number, folder by name out of the tree, file by convention inside it."""
    client = _TreeClient(
        matters=[{"id": "m-ops", "number": "2026-OPS-001"}],
        folders=[],
        files=[
            {
                "id": "tpl",
                "name": "Template - Discovery Set.docx",
                "folder": {"id": "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8"},
                "dateCreated": "2026-08-20",
            }
        ],
        blob=b"PK-firm-template",
    )
    out = resolve_template(client, _CFG, "discovery_set")
    assert isinstance(out, ResolvedTemplate)
    assert out.file_id == "tpl" and out.folder_id == "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8"
    assert out.bytes == b"PK-firm-template"


def test_observed_wire_shape_folder_object_is_read_and_root_files_have_none() -> None:
    from smokeball_connector.library import _entry_folder_id

    assert _entry_folder_id(_LIVE_ENTRY_IN_FOLDER) == "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8"
    assert _entry_folder_id(_LIVE_ENTRY_AT_ROOT) is None
    cfg = LibraryConfig(authored=True, matter_number="2026-OPS-001", folder_name="Document Library", source="t")
    assert is_library_file(_LIVE_ENTRY_IN_FOLDER, cfg, "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8")
    assert not is_library_file(_LIVE_ENTRY_AT_ROOT, cfg, "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8")


def test_resolver_prefers_the_folder_copy_using_the_observed_folder_object_shape() -> None:
    root_copy = dict(_LIVE_ENTRY_IN_FOLDER, id="root", name="Template - Demand Letter.docx", dateCreated="2026-09-01")
    root_copy.pop("folder")
    in_folder = dict(_LIVE_ENTRY_IN_FOLDER, id="lib", name="Template - Demand Letter.docx", dateCreated="2026-01-01")
    client = _FakeClient(
        matters=[{"id": "m-ops", "number": "2026-OPS-001"}],
        folders=[{"id": "9898f74a-3ad9-4b79-b209-a2f0f0c3d7d8", "name": "Document Library"}],
        files=[root_copy, in_folder],
    )
    out = resolve_template(client, _CFG, "demand_letter")
    assert isinstance(out, ResolvedTemplate) and out.file_id == "lib"


def test_library_files_are_recognized_by_name_and_by_folder() -> None:
    assert is_library_file({"name": "Template - Memo.docx"}, _CFG, None)
    assert is_library_file({"name": "template - anything.docx"}, _CFG, None)
    assert is_library_file({"name": "Letterhead.docx", "folderId": "f-lib"}, _CFG, "f-lib")
    assert not is_library_file({"name": "Police Report.pdf", "folderId": "f-docs"}, _CFG, "f-lib")


# ---- the tools ---------------------------------------------------------------------------


def _mock_client(handler) -> object:
    from smokeball_connector.client import SmokeballClient

    client = SmokeballClient(region="us", environment="staging", client_id="cid", client_secret="sec", api_key="apikey")
    client._http = httpx.Client(transport=httpx.MockTransport(handler))
    return client


def _handler(captured: list[httpx.Request], *, template: bytes | None = None, listing_files: list | None = None):
    def handler(request: httpx.Request) -> httpx.Response:
        captured.append(request)
        path = request.url.path
        if path.endswith("/oauth2/token"):
            return httpx.Response(200, json={"access_token": "tok", "expires_in": 3600, "token_type": "Bearer"})
        if request.method == "GET" and path == "/matters":
            return httpx.Response(200, json={"value": [{"id": "m-ops", "number": "2026-OPS-001"}]})
        if request.method == "GET" and path.endswith("/documents/folders"):
            return httpx.Response(200, json={"value": [{"id": "f-lib", "name": "Document Library"}]})
        if request.method == "GET" and path.endswith("/documents/files"):
            return httpx.Response(200, json={"value": listing_files or []})
        if request.method == "GET" and path.endswith("/download"):
            return httpx.Response(200, json={"downloadUrl": _DOWNLOAD_URL, "name": "Template - Memo.docx", "fileExtension": ".docx", "sizeBytes": len(template or b"")})
        if str(request.url) == _DOWNLOAD_URL:
            return httpx.Response(200, content=template or b"")
        if request.method == "POST" and path.endswith("/documents/files"):
            return httpx.Response(202, json={"fileId": "file-88", "uploadUrl": _UPLOAD_URL})
        if str(request.url) == _UPLOAD_URL:
            return httpx.Response(200)
        return httpx.Response(200, json={"ok": True})

    return handler


def _stub_record_check(monkeypatch) -> None:
    from smokeball_connector import record_check as rc

    monkeypatch.setattr(server, "_collect_matter_sources", lambda _m: ([("Src", "text")], [], []))
    monkeypatch.setattr(rc, "run_record_check", lambda *a, **k: rc.RecordCheckResult(passed=True, disposition="pass", refusals=[], checked_sources=1))


def _authored(tmp_path, monkeypatch) -> None:
    path = _write_yaml(tmp_path, "self_initiation:\n  document_library:\n    matter_number: '2026-OPS-001'\n    folder_name: 'Document Library'\n")
    monkeypatch.setenv(CUSTOMER_YAML_ENV, path)


def _put_bytes(captured: list[httpx.Request]) -> bytes:
    return next(r for r in captured if r.method == "PUT").content


def _docx_members(data: bytes) -> dict[str, bytes]:
    """Decompressed zip members by name, for content comparison of two renders.

    NEVER compare two independent renders as raw archive bytes: a .docx is a
    zip whose entry headers carry the current time at 2-second DOS granularity,
    so two renders of identical content that straddle a tick differ only in
    those mtime fields (observed 2026-08-31 as one-run flakes on PRs #2651 and
    #2654; the byte diff was the DOS-time fields in the local headers).
    Comparing member names + decompressed member bytes keeps the full
    byte-for-byte content claim and drops only the timestamp metadata.
    """
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return {name: zf.read(name) for name in zf.namelist()}


def test_no_document_class_is_the_stock_render_byte_for_byte(monkeypatch, tmp_path) -> None:
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", "Body text.")
    assert "formatApplied" not in out
    assert _docx_members(_put_bytes(captured)) == _docx_members(render_markdown_to_docx("Body text."))


def test_document_class_without_an_authored_library_renders_on_the_starter_and_says_so(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv(CUSTOMER_YAML_ENV, str(tmp_path / "absent.yaml"))
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", DISCOVERY_MD, document_class="discovery_set")
    fa = out["formatApplied"]
    assert out["fileId"] == "file-88" and out["refusals"] == []
    assert fa["templateUsed"] is None and fa["templateExpected"] is False
    assert any("not authored" in n for n in fa["notes"])
    assert fa["blocksStyled"]["labels"] == 2
    # The falsifier: class changes the CONTENT. Compared as members, not raw
    # archive bytes — raw bytes differ on zip mtimes alone (see _docx_members),
    # which would let this pass even if the class changed nothing.
    assert _docx_members(_put_bytes(captured)) != _docx_members(render_markdown_to_docx(DISCOVERY_MD))
    assert hashlib.sha256(_put_bytes(captured)).hexdigest() == out["sha256"]


def test_authored_library_but_no_template_file_is_expected_and_loud(monkeypatch, tmp_path) -> None:
    _authored(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured, listing_files=[])))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", "Body.", document_class="memo")
    fa = out["formatApplied"]
    assert out["fileId"] == "file-88"
    assert fa["templateExpected"] is True and fa["templateUsed"] is None
    assert any("Template - Memo.docx" in n for n in fa["notes"])


def test_firm_template_resolves_and_the_draft_renders_into_it(monkeypatch, tmp_path) -> None:
    _authored(tmp_path, monkeypatch)
    template = make_firm_template(header_text="ACME LAW, LLP")
    listing = [{"id": "tpl-1", "name": "Template - Memo.docx", "folderId": "f-lib"}]
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured, template=template, listing_files=listing)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", "Body.", document_class="memo")
    fa = out["formatApplied"]
    assert fa["templateUsed"] == {"name": "Template - Memo.docx", "fileId": "tpl-1", "sha256": hashlib.sha256(template).hexdigest()}
    assert "ACME LAW, LLP" in fa["baseHeaderFooterText"]
    doc = Document(io.BytesIO(_put_bytes(captured)))
    assert doc.sections[0].header.paragraphs[0].text.startswith("ACME LAW, LLP")
    assert [p.text for p in doc.paragraphs if p.text.strip()] == ["Body."]


def test_multi_section_firm_template_is_refused_and_nothing_uploads(monkeypatch, tmp_path) -> None:
    _authored(tmp_path, monkeypatch)
    template = make_firm_template(sections=2)
    listing = [{"id": "tpl-1", "name": "Template - Memo.docx", "folderId": "f-lib"}]
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured, template=template, listing_files=listing)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", "Body.", document_class="memo")
    assert out["fileId"] is None
    assert out["refusals"] and "2 sections" in out["refusals"][0]
    assert not [r for r in captured if r.method == "PUT"]


def test_unknown_document_class_is_a_refusal_with_the_list(monkeypatch, tmp_path) -> None:
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", "Body.", document_class="pleading")
    assert out["fileId"] is None and "discovery_set" in out["refusals"][0]
    assert not [r for r in captured if r.method == "PUT"]


def test_render_docx_template_with_a_class_renders_the_skeleton_on_the_starter(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv(CUSTOMER_YAML_ENV, str(tmp_path / "absent.yaml"))
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    skeleton = "# SKELETON\n## I. Section\n{{FILL: thing | source}}\n"
    out = server.render_docx_template("m-ops", "Template - Memo", skeleton, folder_id="f-lib", document_class="memo")
    assert out["fileId"] == "file-88" and out["fileName"] == "Template - Memo.docx"
    assert out["formatApplied"]["class"] == "memo"
    doc = Document(io.BytesIO(_put_bytes(captured)))
    assert "SMD Item Label" in [s.name for s in doc.styles]  # the starter carries the named styles for the firm to edit
    assert any("{{FILL: thing | source}}" in p.text for p in doc.paragraphs)


def test_template_tool_required_params_are_unchanged() -> None:
    import inspect

    sig = inspect.signature(server.render_docx_template)
    assert [n for n, p in sig.parameters.items() if p.default is inspect.Parameter.empty] == ["matter_id", "file_name", "skeleton_markdown"]
    sig = inspect.signature(server.render_docx_draft)
    assert [n for n, p in sig.parameters.items() if p.default is inspect.Parameter.empty] == ["matter_id", "file_name", "draft_markdown"]


def test_collect_matter_sources_skips_library_templates(monkeypatch, tmp_path) -> None:
    """Templates are not record: a header-only letterhead extracts to nothing
    and would otherwise refuse the whole record check."""
    _authored(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    listing = [
        {"id": "tpl", "name": "Template - Memo.docx", "folderId": "f-lib"},
        {"id": "lh", "name": "Letterhead.docx", "folderId": "f-lib"},
        {"id": "rec", "name": "Police Report.txt", "folderId": "f-docs"},
    ]

    def handler(request: httpx.Request) -> httpx.Response:
        captured.append(request)
        path = request.url.path
        if path.endswith("/oauth2/token"):
            return httpx.Response(200, json={"access_token": "tok", "expires_in": 3600, "token_type": "Bearer"})
        if request.method == "GET" and path.endswith("/documents/folders"):
            return httpx.Response(200, json={"value": [{"id": "f-lib", "name": "Document Library"}]})
        if request.method == "GET" and path.endswith("/documents/files"):
            return httpx.Response(200, json={"value": listing})
        if request.method == "GET" and path.endswith("/download"):
            return httpx.Response(200, json={"downloadUrl": _DOWNLOAD_URL, "name": "Police Report.txt", "fileExtension": ".txt", "sizeBytes": 11})
        if str(request.url) == _DOWNLOAD_URL:
            return httpx.Response(200, content=b"REPORT TEXT")
        return httpx.Response(200, json={"ok": True})

    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(handler))
    sources, _vision, unextractable = server._collect_matter_sources("m-ops")
    assert sources == [("Police Report.txt", "REPORT TEXT")]
    assert unextractable == []
    downloaded = [r.url.path for r in captured if r.url.path.endswith("/download")]
    assert downloaded and all("/rec/" in d for d in downloaded)


@pytest.mark.parametrize("cls", ["discovery_set", "mediation_brief", "letter"])
def test_every_class_files_on_the_starter(monkeypatch, tmp_path, cls: str) -> None:
    monkeypatch.setenv(CUSTOMER_YAML_ENV, str(tmp_path / "absent.yaml"))
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "Draft", DISCOVERY_MD, document_class=cls)
    assert out["fileId"] == "file-88" and out["formatApplied"]["class"] == cls


# --- One naming authority (ss#2490) -------------------------------------------
#
# The renderer resolves a class's template BY NAME, so a template filed under any
# other name is inert while the delivery note still reads as if the format were
# live. Found on pilot 2026-08-20: three templates filed, one live
# (vfy_01M0GHR4T1V6HG46BVX3S61BCC). The skill body said "name it by the
# convention unless the blessing named it otherwise", which assumed no override
# existed yet — an ordering assumption that is false on any seat configured in an
# earlier cycle. Prose could not close it; these pin the mechanical check.


def _authored_with_override(tmp_path, monkeypatch) -> None:
    path = _write_yaml(
        tmp_path,
        "self_initiation:\n"
        "  document_library:\n"
        "    matter_number: '2026-OPS-001'\n"
        "    folder_name: 'Document Library'\n"
        "    templates:\n"
        "      demand_letter: 'Firm Demand Shell'\n",
    )
    monkeypatch.setenv(CUSTOMER_YAML_ENV, path)


def test_names_agree_ignores_case_space_and_extension() -> None:
    from smokeball_connector.library import names_agree

    assert names_agree("Template - Letter.docx", "Template - Letter")
    assert names_agree(" template - letter ", "Template - Letter.docx")
    assert not names_agree("Template - Letter.docx", "Template - Demand Letter.docx")


def test_class_template_name_is_reported_so_the_skill_need_not_guess(monkeypatch, tmp_path) -> None:
    _authored_with_override(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    out = server.render_docx_template("m-1", "Firm Demand Shell.docx", "# Shell\n", document_class="demand_letter")
    assert out["formatApplied"]["classTemplateName"] == "Firm Demand Shell.docx"
    assert out["refusals"] == []


def test_filing_a_class_template_under_a_name_the_renderer_will_not_open_is_refused(monkeypatch, tmp_path) -> None:
    """THE ss#2490 REGRESSION. The convention name looks right and is wrong here,
    because this seat authors an override. Refused, not filed-and-warned."""
    _authored_with_override(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    out = server.render_docx_template("m-1", "Template - Demand Letter.docx", "# Shell\n", document_class="demand_letter")
    assert out["fileId"] is None
    assert out["refusals"] and "Firm Demand Shell.docx" in out["refusals"][0]
    # The falsifier that matters: NOTHING was uploaded. A warning-only fix would
    # still have PUT the bytes, which is exactly the state that shipped.
    assert not [r for r in captured if r.method == "PUT"]


def test_the_convention_name_is_accepted_when_no_override_is_authored(monkeypatch, tmp_path) -> None:
    _authored(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    out = server.render_docx_template("m-1", "Template - Demand Letter.docx", "# Shell\n", document_class="demand_letter")
    assert out["refusals"] == [] and out["fileId"] == "file-88"


def test_a_draft_is_never_subject_to_the_class_template_name(monkeypatch, tmp_path) -> None:
    """The guard is for templates only. A DRAFT is named for the matter and must
    keep its own name, or every drafting lane breaks."""
    _authored_with_override(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    _stub_record_check(monkeypatch)
    out = server.render_docx_draft("m-1", "DRAFT Demand - Alvarez", DISCOVERY_MD, document_class="demand_letter")
    assert out["refusals"] == [] and out["fileId"] == "file-88"


def test_no_document_class_means_no_name_opinion(monkeypatch, tmp_path) -> None:
    _authored_with_override(tmp_path, monkeypatch)
    captured: list[httpx.Request] = []
    monkeypatch.setattr(server, "_get_client", lambda: _mock_client(_handler(captured)))
    out = server.render_docx_template("m-1", "Anything At All.docx", "# Shell\n")
    assert out["refusals"] == [] and "formatApplied" not in out

# ---- the three-state matter lookup (ss-console#2536) ---------------------------------
#
# ``find_matter_id`` answers "the id, or None", which is right for template
# resolution: a draft still files on the starter and the note says the template
# did not resolve. It is WRONG for deduping a create, where None conflates "there
# is no such matter" with "I could not look", and only one of those makes it safe
# to write. These pin the third state.


class _ListClient:
    def __init__(self, matters, *, list_error=None, search_error=None) -> None:
        self.matters, self.list_error, self.search_error = matters, list_error, search_error

    def get(self, path: str, **params):
        assert path == "/matters"
        if "Search" in params:
            if self.search_error:
                raise self.search_error
            return {"value": self.matters}
        if self.list_error:
            raise self.list_error
        return {"value": self.matters}


def test_lookup_reports_found_with_what_it_matched_on() -> None:
    client = _ListClient([{"id": "m-9", "number": "OPS-1"}])
    result = lookup_matter(client, number="OPS-1")
    assert (result.state, result.matter_id, result.matched_on) == (LOOKUP_FOUND, "m-9", "number")
    assert result.found is True


def test_lookup_reports_not_found_when_the_enumeration_completed() -> None:
    result = lookup_matter(_ListClient([{"id": "m-9", "number": "OTHER"}]), number="OPS-1")
    assert result.state == LOOKUP_NOT_FOUND
    assert result.found is False and result.failed is False


def test_an_incomplete_enumeration_is_a_failure_and_not_an_empty_answer() -> None:
    result = lookup_matter(_ListClient([], list_error=RuntimeError("boom")), number="OPS-1")
    assert result.state == LOOKUP_FAILED
    assert result.failed is True
    assert "could not be read" in result.reason


def test_lookup_matches_the_description_and_client_pair_too() -> None:
    client = _ListClient(
        [{"id": "m-4", "number": "OTHER", "description": "Operator Library", "clientIds": ["c-1"]}]
    )
    result = lookup_matter(
        client, number="OPS-1", description="Operator Library", client_contact_id="c-1"
    )
    assert result.matter_id == "m-4"
    assert result.matched_on == "description and client"


def test_the_description_pair_needs_both_halves() -> None:
    client = _ListClient(
        [{"id": "m-4", "number": "OTHER", "description": "Operator Library", "clientIds": ["c-2"]}]
    )
    result = lookup_matter(
        client, number="OPS-1", description="Operator Library", client_contact_id="c-1"
    )
    assert result.state == LOOKUP_NOT_FOUND


def test_client_ids_are_read_across_the_shapes_the_api_uses() -> None:
    for shape in (
        {"clientIds": ["c-1"]},
        {"clientIds": [{"id": "c-1"}]},
        {"clients": [{"id": "c-1"}]},
        {"clientId": "c-1"},
    ):
        matter = {"id": "m-4", "number": "OTHER", "description": "Operator Library", **shape}
        result = lookup_matter(
            _ListClient([matter]),
            number="OPS-1",
            description="Operator Library",
            client_contact_id="c-1",
        )
        assert result.found is True, shape


def test_find_matter_id_keeps_its_two_state_answer_for_template_resolution() -> None:
    """The old signature and the old semantics, unchanged: resolution treats a
    failed lookup and an absent matter the same way, because both end in a draft
    on the starter base with the reason reported."""
    from smokeball_connector.library import find_matter_id

    assert find_matter_id(_ListClient([{"id": "m-9", "number": "OPS-1"}]), "OPS-1") == "m-9"
    assert find_matter_id(_ListClient([]), "OPS-1") is None
    assert find_matter_id(_ListClient([], list_error=RuntimeError("boom")), "OPS-1") is None
