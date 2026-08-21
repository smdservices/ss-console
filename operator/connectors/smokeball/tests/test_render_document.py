"""docx_format.render_document: code owns typography, the firm's template wins.

Assertions run at two layers on purpose: python-docx's object model for what it
models (styles, paragraph formats, sections) and the raw XML parts for what
Word reads that python-docx does not model (table borders, the PAGE field,
docDefaults, the four rFonts attributes). Nothing here opens Word; the Word
open is a runtime row of the contract, performed by a person.

The "firm template" is BUILT in-test (no checked-in binaries: a .docx in a PR
is unreviewable and rots). It deliberately carries the pathologies of real
firm files the spike found: no ``Heading 1`` / ``List Bullet`` / ``Table Grid``
styles, a header with an image, non-default margins, an optional second
section, an optional .dotx content type.
"""

from __future__ import annotations

import io
import struct
import zipfile
import zlib

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt

from smokeball_connector.docx_format import (
    CLASS_RULES,
    DOCUMENT_CLASSES,
    FormatRefused,
    FormatReport,
    render_document,
)

_DOCX_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
_DOTX_MAIN = "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"


def _png_1x1() -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = struct.pack(">I", len(data)) + tag + data
        return c + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = zlib.compress(b"\x00\xff\x00\x00")
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", raw) + chunk(b"IEND", b"")


def make_firm_template(
    *,
    header_text: str = "ACME LAW, LLP",
    header_image: bool = True,
    margins_in: float = 1.25,
    named_styles: dict[str, dict] | None = None,
    footer_text: str = "",
    drop_styles: tuple[str, ...] = ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "Table Grid"),
    body_text: str = "OLD TEMPLATE BODY that must not survive",
    sections: int = 1,
    dotx: bool = False,
) -> bytes:
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(margins_in)
    hdr = doc.sections[0].header
    hp = hdr.paragraphs[0]
    hp.text = header_text
    if header_image:
        hp.add_run().add_picture(io.BytesIO(_png_1x1()), width=Inches(0.2))
    if footer_text:
        doc.sections[0].footer.paragraphs[0].text = footer_text
    doc.add_paragraph(body_text)
    for extra in range(sections - 1):
        doc.add_section()
        doc.add_paragraph(f"section {extra + 2}")
    for name, spec in (named_styles or {}).items():
        st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        st.font.name = spec.get("font", "Arial")
        st.font.size = Pt(spec.get("size", 14))
        if spec.get("indent"):
            st.paragraph_format.first_line_indent = Inches(spec["indent"])
    for name in drop_styles:
        try:
            doc.styles[name].element.getparent().remove(doc.styles[name].element)
        except KeyError:
            pass
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if dotx:
        src = zipfile.ZipFile(io.BytesIO(data))
        out = io.BytesIO()
        with zipfile.ZipFile(out, "w") as dst:
            for item in src.infolist():
                payload = src.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    payload = payload.replace(_DOCX_MAIN.encode(), _DOTX_MAIN.encode())
                dst.writestr(item, payload)
        data = out.getvalue()
    return data


DISCOVERY_MD = """| PLAINTIFF JANE DOE, | Case No. {{FILL: case number | operative pleading}} |
| v. |  |
| DEFENDANT ACME CORP. |  |

# PLAINTIFF'S SPECIAL INTERROGATORIES TO DEFENDANT, SET ONE

## DEFINITIONS
"INCIDENT" means the collision described in the complaint.

**SPECIAL INTERROGATORY NO. 1:**
Identify each person who witnessed the INCIDENT.

**SPECIAL INTERROGATORY NO. 2:**
State all facts supporting your denial. {{NOT IN RECORD: prior denial}}

| Provider | Balance |
| --- | --- |
| Dr. A | $1,200.00 |

SPECIAL INTERROGATORY NO. 1 asks about witnesses, and this long prose sentence is body text because a label is a short line and this is not one of those.
"""


def _parts(data: bytes) -> dict[str, str]:
    z = zipfile.ZipFile(io.BytesIO(data))
    return {n: z.read(n).decode("utf-8", "replace") for n in z.namelist() if n.endswith(".xml")}


def _styled(data: bytes) -> list[tuple[str, str]]:
    doc = Document(io.BytesIO(data))
    return [(p.style.name, p.text[:28]) for p in doc.paragraphs if p.text.strip()]


# ---- starter base ------------------------------------------------------------


def test_starter_base_rewrites_doc_defaults_to_the_authored_standard() -> None:
    data, _ = render_document("Body.", "memo", None)
    styles = _parts(data)["word/styles.xml"]
    assert 'w:ascii="Times New Roman"' in styles and 'w:eastAsia="Times New Roman"' in styles
    assert '<w:sz w:val="24"/>' in styles
    assert 'w:line="240"' in styles


def test_starter_base_defines_every_named_style_and_honors_them() -> None:
    data, report = render_document(DISCOVERY_MD, "discovery_set", None)
    r = report.to_dict()
    assert r["fallbacks"] == []
    assert {"SMD Item Label", "SMD Item Text", "SMD Heading 1", "SMD Heading 2", "SMD Body", "SMD Caption"} <= set(r["stylesHonored"])
    assert r["blocksStyled"] == {"labels": 2, "tables": 2, "headings": 2}  # "## DEFINITIONS" is a heading


def test_golden_discovery_set_sequence() -> None:
    data, _ = render_document(DISCOVERY_MD, "discovery_set", None)
    seq = _styled(data)
    expected = [
        ("SMD Heading 1", "PLAINTIFF'S SPECIAL"),
        ("SMD Heading 2", "DEFINITIONS"),
        ("SMD Body", '"INCIDENT" means'),
        ("SMD Item Label", "SPECIAL INTERROGATORY NO. 1:"),
        ("SMD Item Text", "Identify each person"),
        ("SMD Item Label", "SPECIAL INTERROGATORY NO. 2:"),
        ("SMD Item Text", "State all facts"),
    ]
    assert [(s, t.startswith(p)) for (s, t), (_, p) in zip(seq, expected)] == [(s, True) for s, _ in expected]
    assert seq[-1][0] == "SMD Body"  # the long prose line is body, not a label


def test_markers_survive_verbatim_inside_table_cells_and_items() -> None:
    data, _ = render_document(DISCOVERY_MD, "discovery_set", None)
    doc = Document(io.BytesIO(data))
    caption = doc.tables[0]
    assert caption.cell(0, 1).text == "Case No. {{FILL: case number | operative pleading}}"
    assert any("{{NOT IN RECORD: prior denial}}" in p.text for p in doc.paragraphs)


def test_caption_table_gets_only_the_inside_vertical_rule_and_data_tables_get_a_grid() -> None:
    data, _ = render_document(DISCOVERY_MD, "discovery_set", None)
    xml = _parts(data)["word/document.xml"]
    borders = xml.split("<w:tblBorders>")
    assert len(borders) == 3  # two tables
    caption, data_table = borders[1].split("</w:tblBorders>")[0], borders[2].split("</w:tblBorders>")[0]
    assert "<w:insideV" in caption and "<w:top" not in caption
    assert all(f"<w:{e}" in data_table for e in ("top", "left", "bottom", "right", "insideH", "insideV"))


def test_letter_class_first_table_is_a_plain_grid() -> None:
    data, _ = render_document("| a | b |\n| c | d |", "letter", None)
    xml = _parts(data)["word/document.xml"]
    assert "<w:top" in xml.split("<w:tblBorders>")[1]


def test_page_number_field_added_when_the_base_has_no_footer() -> None:
    data, report = render_document("Body.", "memo", None)
    parts = _parts(data)
    footer = next(v for k, v in parts.items() if k.startswith("word/footer"))
    assert 'w:fldCharType="begin"' in footer and "PAGE" in footer
    assert "page number field added to the footer" in report.to_dict()["notes"]


def test_letter_class_adds_no_page_number_by_default() -> None:
    data, _ = render_document("Body.", "letter", None)
    parts = _parts(data)
    assert not any("PAGE" in v for k, v in parts.items() if k.startswith("word/footer"))


@pytest.mark.parametrize("cls", DOCUMENT_CLASSES)
def test_every_class_renders_the_same_content_without_error(cls: str) -> None:
    data, report = render_document(DISCOVERY_MD, cls, None)
    assert data[:2] == b"PK"
    assert report.to_dict()["class"] == cls


def test_unknown_class_is_a_value_error() -> None:
    with pytest.raises(ValueError):
        render_document("x", "pleading_paper", None)


def test_mediation_brief_headings_centered_bold_roman_then_indented_underlined_letters() -> None:
    data, _ = render_document("# I. INTRODUCTION\n## A. Parties\n### 1. Plaintiff\nBody.", "mediation_brief", None)
    doc = Document(io.BytesIO(data))
    h1, h2, h3 = doc.paragraphs[0], doc.paragraphs[1], doc.paragraphs[2]
    assert h1.text == "I. INTRODUCTION" and h1.style.name == "SMD Heading 1"
    assert h2.text == "A. Parties" and h2.style.name == "SMD Heading 2"
    assert h3.style.name == "SMD Heading 3"
    # the numerals are CONTENT: nothing was added or renumbered
    assert [p.text for p in doc.paragraphs[:3]] == ["I. INTRODUCTION", "A. Parties", "1. Plaintiff"]


# ---- firm-supplied base (provenance a) ------------------------------------------


def test_firm_template_header_image_margins_and_footer_survive_and_body_is_cleared() -> None:
    base = make_firm_template(header_text="ACME LAW, LLP", footer_text="123 Main St.")
    data, report = render_document(DISCOVERY_MD, "discovery_set", base)
    doc = Document(io.BytesIO(data))
    assert doc.sections[0].header.paragraphs[0].text.startswith("ACME LAW, LLP")
    assert doc.sections[0].left_margin == Inches(1.25)
    assert any(n.startswith("word/media/") for n in zipfile.ZipFile(io.BytesIO(data)).namelist())
    assert "OLD TEMPLATE BODY" not in "\n".join(p.text for p in doc.paragraphs)
    r = report.to_dict()
    assert "ACME LAW, LLP" in r["baseHeaderFooterText"] and "123 Main St." in r["baseHeaderFooterText"]
    # the firm's own footer is left alone: no PAGE field injected over it
    assert "page number field added to the footer" not in r["notes"]


def test_firm_template_lacking_stock_styles_does_not_raise_and_falls_back_inline() -> None:
    base = make_firm_template()  # Heading 1-3 / List Bullet / Table Grid removed
    data, report = render_document(DISCOVERY_MD + "- a bullet\n", "discovery_set", base)
    r = report.to_dict()
    assert set(r["fallbacks"]) >= {"SMD Item Label", "SMD Item Text", "SMD Heading 1", "SMD Body"}
    doc = Document(io.BytesIO(data))
    label = next(p for p in doc.paragraphs if p.text == "SPECIAL INTERROGATORY NO. 1:")
    assert all(run.bold and run.underline and run.font.all_caps for run in label.runs if run.text.strip())
    item = next(p for p in doc.paragraphs if p.text.startswith("Identify each person"))
    assert item.paragraph_format.first_line_indent == Inches(0.5)
    assert item.paragraph_format.line_spacing == 2.0
    bullet = next(p for p in doc.paragraphs if "a bullet" in p.text)
    assert bullet.text.startswith("•")


def test_firm_template_that_defines_a_named_style_wins_and_is_not_a_fallback() -> None:
    base = make_firm_template(named_styles={"SMD Item Label": {"font": "Arial", "size": 14}})
    data, report = render_document(DISCOVERY_MD, "discovery_set", base)
    r = report.to_dict()
    assert "SMD Item Label" in r["stylesHonored"] and "SMD Item Label" not in r["fallbacks"]
    doc = Document(io.BytesIO(data))
    label = next(p for p in doc.paragraphs if p.text == "SPECIAL INTERROGATORY NO. 1:")
    assert label.style.name == "SMD Item Label"
    assert doc.styles["SMD Item Label"].font.name == "Arial"  # the firm's definition, untouched


def test_firm_styles_are_never_written_back() -> None:
    base = make_firm_template()
    data, _ = render_document("Body.", "memo", base)
    styles = _parts(data)["word/styles.xml"]
    assert "SMD Body" not in styles and "SMD Item Label" not in styles


def test_dotx_base_opens_and_is_noted() -> None:
    base = make_firm_template(dotx=True)
    data, report = render_document("Body.", "memo", base)
    assert data[:2] == b"PK"
    assert any("dotx" in n for n in report.to_dict()["notes"])
    assert _DOCX_MAIN in _parts(data)["[Content_Types].xml"]


def test_multi_section_base_is_refused_and_nothing_is_rendered() -> None:
    base = make_firm_template(sections=2)
    with pytest.raises(FormatRefused, match="2 sections"):
        render_document("Body.", "memo", base)


def test_report_round_trips_template_used_and_expected_flags() -> None:
    report = FormatReport(document_class="memo", template_used={"name": "Template - Memo.docx", "fileId": "f1", "sha256": "abc"}, template_expected=True)
    _, out = render_document("Body.", "memo", None, report)
    d = out.to_dict()
    assert d["templateUsed"]["name"] == "Template - Memo.docx" and d["templateExpected"] is True


def test_class_rules_cover_every_class() -> None:
    assert set(CLASS_RULES) == set(DOCUMENT_CLASSES)


# --- Heading delegation to the firm's own styles (ss#2448) ---------------------
#
# The promise is "the firm edits its Word template and the next draft follows".
# For a firm template that never heard of our SMD names, the shipped renderer
# formatted headings INLINE, so the typography was frozen into each draft and a
# later template edit moved nothing (observed on pilot,
# vfy_01M0GK1SS4CZKGHMV8R1CFAZSG). Delegating the heading roles to Word's own
# Heading 1-3 makes the promise true by construction.
#
# Deliberately narrow. Body and item text are NOT delegated: a plain paragraph
# already resolves to Normal, so font follows the firm today, and the inline
# branch is carrying the class's court-required LAYOUT.


def test_a_heading_delegates_to_the_bases_own_heading_style() -> None:
    """The base defines Heading 2 but not SMD Heading 2: use theirs."""
    base = make_firm_template(drop_styles=("List Bullet", "Table Grid"))  # keeps Heading 1-3
    report = FormatReport(document_class="mediation_brief")
    blob, report = render_document("## II. Argument\n\nText.\n", "mediation_brief", base, report)
    d = report.to_dict()
    assert d["stylesDelegated"] == {"SMD Heading 2": "Heading 2"}
    assert "SMD Heading 2" not in d["fallbacks"]
    doc = Document(io.BytesIO(blob))
    para = next(p for p in doc.paragraphs if "Argument" in p.text)
    assert para.style.name == "Heading 2"


def test_a_delegated_heading_still_carries_the_classs_required_layout() -> None:
    """The firm's style supplies font; the COURT expects a centered level-1
    heading on a mediation brief and a firm's built-in Heading 1 is left."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    base = make_firm_template(drop_styles=("List Bullet", "Table Grid"))
    report = FormatReport(document_class="mediation_brief")
    blob, _ = render_document("# I. Introduction\n\nText.\n", "mediation_brief", base, report)
    doc = Document(io.BytesIO(blob))
    para = next(p for p in doc.paragraphs if "Introduction" in p.text)
    assert para.style.name == "Heading 1"
    assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER
    # ...and emphasis is left to the firm's style rather than forced over it.
    assert not any(r.bold for r in para.runs)


def test_a_base_defining_neither_still_goes_inline_and_says_what_to_add() -> None:
    base = make_firm_template()  # drops Heading 1-3 by default
    report = FormatReport(document_class="mediation_brief")
    _blob, report = render_document("## II. Argument\n\nText.\n", "mediation_brief", base, report)
    d = report.to_dict()
    assert d["stylesDelegated"] == {}
    assert "SMD Heading 2" in d["fallbacks"]
    assert any("add 'Heading 2' in Word" in n for n in d["notes"])


def test_item_text_is_never_delegated_so_a_served_set_keeps_its_double_spacing() -> None:
    """THE REGRESSION A WIDER LADDER WOULD HAVE SHIPPED. Delegating SMD Item
    Text to Normal suppresses the inline branch that carries the authored
    double-spacing between requests, and a discovery set is SERVED."""
    base = make_firm_template(drop_styles=("List Bullet", "Table Grid"))
    report = FormatReport(document_class="discovery_set")
    blob, report = render_document(DISCOVERY_MD, "discovery_set", base, report)
    assert "SMD Item Text" not in report.to_dict()["stylesDelegated"]
    doc = Document(io.BytesIO(blob))
    answers = [p for p in doc.paragraphs if p.text.startswith("Identify each person")]
    assert answers, "expected the item text paragraph"
    assert answers[0].paragraph_format.line_spacing == 2.0


def test_a_character_style_named_like_a_heading_does_not_kill_the_render() -> None:
    """python-docx raises ValueError (not KeyError) assigning a character style
    to a paragraph, and has_style catches only KeyError."""
    doc = Document()
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[name].element.getparent().remove(doc.styles[name].element)
    doc.styles.add_style("Heading 2", WD_STYLE_TYPE.CHARACTER)
    buf = io.BytesIO()
    doc.save(buf)
    report = FormatReport(document_class="mediation_brief")
    blob, report = render_document("## II. Argument\n\nText.\n", "mediation_brief", buf.getvalue(), report)
    d = report.to_dict()
    assert d["stylesDelegated"] == {}
    assert "SMD Heading 2" in d["fallbacks"]
    assert blob


def test_an_outline_numbered_heading_is_not_delegated_to() -> None:
    """The model writes the numeral itself ("## I. Introduction"), so a firm's
    multilevel-list Heading 2 would render "1.1 I. Introduction"."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Heading 2"]
    ppr = style.element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    buf = io.BytesIO()
    doc.save(buf)
    report = FormatReport(document_class="mediation_brief")
    _blob, report = render_document("## II. Argument\n\nText.\n", "mediation_brief", buf.getvalue(), report)
    d = report.to_dict()
    assert d["stylesDelegated"] == {}, "a numbered base style must not be delegated to"
    assert "SMD Heading 2" in d["fallbacks"]


def test_numbering_inherited_through_the_base_style_chain_is_also_refused() -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    parent = doc.styles.add_style("Numbered Parent", WD_STYLE_TYPE.PARAGRAPH)
    ppr = parent.element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    ppr.append(num_pr)
    doc.styles["Heading 2"].base_style = parent
    buf = io.BytesIO()
    doc.save(buf)
    report = FormatReport(document_class="mediation_brief")
    _blob, report = render_document("## II. Argument\n\nText.\n", "mediation_brief", buf.getvalue(), report)
    assert report.to_dict()["stylesDelegated"] == {}


def test_our_own_named_style_still_wins_over_the_bases_equivalent() -> None:
    base = make_firm_template(
        drop_styles=("List Bullet", "Table Grid"),
        named_styles={"SMD Heading 2": {"font": "Courier New", "size": 15}},
    )
    report = FormatReport(document_class="mediation_brief")
    blob, report = render_document("## II. Argument\n\nText.\n", "mediation_brief", base, report)
    d = report.to_dict()
    assert d["stylesDelegated"] == {} and "SMD Heading 2" in d["stylesHonored"]
    doc = Document(io.BytesIO(blob))
    assert next(p for p in doc.paragraphs if "Argument" in p.text).style.name == "SMD Heading 2"
